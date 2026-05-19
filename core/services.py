import re
import docx
from .models import Office, StrategicLevel, Indicator, PerformanceRecord

_WORD_TO_QUARTER = {
    "first": 1,
    "second": 2,
    "third": 3,
    "fourth": 4,
}

_QUARTER_PATTERNS = [
    re.compile(r"\b([1-4])\s*(?:st|nd|rd|th)\s+Quarter\b", re.IGNORECASE),
    re.compile(r"\b(First|Second|Third|Fourth)\s+Quarter\b", re.IGNORECASE),
    re.compile(r"\bQuarter\s*([1-4])\b", re.IGNORECASE),
    re.compile(r"\b([1-4])\s*Quarter\b", re.IGNORECASE),
    re.compile(r"\bQ\s*([1-4])\b", re.IGNORECASE),
    re.compile(r"\b([1-4])\s*Q\b", re.IGNORECASE),
]


def _harvest_period_hint_text(doc):
    """Title-style text only (paragraphs + first two table rows)."""
    parts = []
    for p in doc.paragraphs[:50]:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    if doc.tables:
        table = doc.tables[0]
        for row in table.rows[:2]:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _detect_quarter_year(text):
    """Infer quarter and calendar year from title / filename text.

    Handles ``LPC-OPMM-Q4-2025`` style names. Uses the **rightmost** ``20xx``
    year when no year is paired with the quarter, so an early ``2024`` note
    does not override ``2025`` at the end of the title.
    """
    quarter = None
    year = None
    if not text or not text.strip():
        return quarter, year
    t = text.strip()
    head = t[:600]

    # Strong signals: Q4-2025, Q4_2025, Q4 2025 (filenames and titles)
    m_pair = re.search(r'(?i)\bQ\s*([1-4])\s*[-_.\s]*\s*(20\d{2})\b', head)
    if m_pair:
        quarter = int(m_pair.group(1))
        year = int(m_pair.group(2))
    else:
        mf = re.search(
            r'(?i)\b(First|Second|Third|Fourth)\s+Quarter\s*[-_,.\s]*\s*(20\d{2})\b',
            head,
        )
        if mf:
            qw = _WORD_TO_QUARTER.get(mf.group(1).lower())
            if qw:
                quarter = qw
                year = int(mf.group(2))

    best_pos = len(t) + 1
    if quarter is None:
        for pat in _QUARTER_PATTERNS:
            for m in pat.finditer(t):
                raw = m.group(1)
                if raw is None:
                    continue
                if raw.isdigit():
                    q = int(raw)
                else:
                    q = _WORD_TO_QUARTER.get(raw.lower())
                if q and 1 <= q <= 4 and m.start() < best_pos:
                    best_pos = m.start()
                    quarter = q

    if year is None:
        years = [
            int(m.group(1))
            for m in re.finditer(r'\b(20[0-9]{2})\b', t)
            if 2000 <= int(m.group(1)) <= 2099
        ]
        if years:
            year = years[-1]
            if quarter is not None and best_pos < len(t):
                after = t[best_pos : best_pos + 40]
                m_after = re.search(r'\b(20[0-9]{2})\b', after)
                if m_after:
                    y2 = int(m_after.group(1))
                    if 2000 <= y2 <= 2099:
                        year = y2
    return quarter, year


FY_YEAR_MARKERS = re.compile(
    r'\b(?:FY|SFY|CY)\s*20[0-9]{2}\b|\bFY20[0-9]{2}\b',
    re.IGNORECASE,
)

# "Concerned Office / Campus" cells often list many units: "Chancellor, OVCAF, OVCAA, …"
_CONCERNED_OFFICE_SPLIT = re.compile(r'[,;\n\r]+')


def split_concerned_office_cell(raw):
    """Return distinct office names from one LPC / FYDP cell (comma- or line-separated)."""
    if raw is None:
        return []
    s = str(raw).strip()
    if not s:
        return []
    parts = _CONCERNED_OFFICE_SPLIT.split(s)
    out = []
    seen = set()
    for p in parts:
        name = p.strip()
        if len(name) < 2:
            continue
        key = name.casefold()
        if key not in seen:
            seen.add(key)
            out.append(name)
    return out


def _strip_fiscal_year_markers(text):
    """Remove FY/CY style year phrases so bare ``2025`` is not mistaken for a KPI target."""
    if not text:
        return ''
    t = FY_YEAR_MARKERS.sub(' ', str(text))
    t = re.sub(r'\bfiscal\s+year\s*20[0-9]{2}\b', ' ', t, flags=re.IGNORECASE)
    return re.sub(r'\s+', ' ', t).strip()


def extract_target_for_monitor(indicator_text):
    """Parse planned target from the indicator column.

    Prefer explicit ``(number)`` targets (skipping parenthetical four-digit calendar years).
    Never treat ``FY 2025``-style markers as the numeric target when no real KPI is present.
    """
    if not indicator_text or not str(indicator_text).strip():
        return None
    s = str(indicator_text).strip()
    for token in re.findall(r'\((\d+(?:\.\d+)?)\)', s):
        if not re.fullmatch(r'20\d{2}', token):
            return float(token)
    stripped = _strip_fiscal_year_markers(s)
    stripped = re.sub(r'\(\s*20\d{2}\s*\)', ' ', stripped)
    m = re.search(r'(?:target|goal)\s*[:=]?\s*(\d+(?:\.\d+)?)', stripped, re.I)
    if m:
        return float(m.group(1))
    m = re.search(r'(\d+(?:\.\d+)?)', stripped)
    if not m:
        return None
    token = m.group(1)
    if re.fullmatch(r'20\d{2}', token):
        return None
    return float(token)


def extract_actual_for_monitor(accomplishment_text):
    """Parse accomplishment cell; ``None`` when empty or no usable number."""
    if not accomplishment_text or not str(accomplishment_text).strip():
        return None
    s = str(accomplishment_text).strip()
    for token in re.findall(r'\((\d+(?:\.\d+)?)\)', s):
        if not re.fullmatch(r'20\d{2}', token):
            return float(token)
    stripped = _strip_fiscal_year_markers(s)
    stripped = re.sub(r'\(\s*20\d{2}\s*\)', ' ', stripped)
    m = re.search(r'(\d+(?:\.\d+)?)', stripped)
    if not m:
        return None
    token = m.group(1)
    if re.fullmatch(r'20\d{2}', token):
        return None
    return float(token)


def extract_number(text):
    """Backward-compatible helper (tests): try indicator-style then accomplishment-style."""
    if not text:
        return 0.0
    v = extract_target_for_monitor(text)
    if v is not None:
        return float(v)
    v = extract_actual_for_monitor(text)
    return float(v) if v is not None else 0.0


def _cell_value_str(value):
    if value is None:
        return ''
    return str(value).strip()


def _ws_cell(ws, row, col_1based):
    if col_1based is None:
        return ''
    return _cell_value_str(ws.cell(row=row, column=col_1based).value)


def _normalize_excel_header_text(value):
    s = _cell_value_str(value).lower()
    s = re.sub(r'\s+', ' ', s)
    return s.strip()


def _matrix_label_changed(prev, new):
    """True when ``new`` is a different merged label than ``prev`` (not empty repeat)."""
    if not (new or '').strip():
        return False
    if not (prev or '').strip():
        return True
    return _normalize_excel_header_text(new) != _normalize_excel_header_text(prev)


def _word_cell_above(table, row_idx, col_idx):
    """Text in ``col_idx`` on the row immediately above ``row_idx`` (typical 2-row Word merge)."""
    if row_idx < 1 or col_idx is None or col_idx < 0 or not table or row_idx >= len(table.rows):
        return ''
    prev = _word_table_row_cells_grid(table, row_idx - 1)
    if col_idx >= len(prev):
        return ''
    return (prev[col_idx] or '').strip()


def _word_table_ncols(table):
    """Number of layout-grid columns (max across rows)."""
    try:
        n = len(table.columns)
        if n > 0:
            return n
    except Exception:
        pass
    n = 0
    for row in table.rows[:16]:
        try:
            n = max(n, len(row.cells))
        except Exception:
            continue
    return max(n, 7)


def _word_table_row_cells_grid(table, row_idx):
    """One text per grid column; python-docx already duplicates merged cells across columns.

    Returns the column-aligned text list for ``row_idx``. Horizontally merged cells
    repeat in subsequent columns (so fill-down logic still sees the merged text);
    rows with fewer logical cells than the grid are right-padded with empty strings.
    """
    if row_idx < 0 or row_idx >= len(table.rows):
        return [''] * _word_table_ncols(table)
    try:
        texts = [(c.text or '').strip() for c in table.rows[row_idx].cells]
    except Exception:
        texts = []
    ncols = _word_table_ncols(table)
    if len(texts) < ncols:
        texts = texts + [''] * (ncols - len(texts))
    return texts


def _word_table_vertical_merge_fill(table, start_row_idx, col_idx, data_start, header_row_idx=-1):
    """Last non-empty text in ``col_idx`` scanning upward from ``start_row_idx``.

    Scans through the matrix header row (``data_start - 1``) and, when needed, rows down to
    ``header_row_idx`` so we do not pull random title text from rows above the real column
    headers (that produced duplicate KPI rows in multi-table fixtures).

    When ``header_row_idx < 0`` (legacy layout), the scan floor is ``max(0, data_start - 1)``.
    """
    if col_idx is None or col_idx < 0 or not table:
        return ''
    top = min(int(start_row_idx), len(table.rows) - 1)
    min_r = max(0, int(data_start) - 1)
    if header_row_idx >= 0:
        min_r = max(min_r, int(header_row_idx))
    for r in range(top, min_r - 1, -1):
        row_cells = _word_table_row_cells_grid(table, r)
        if col_idx >= len(row_cells):
            continue
        t = (row_cells[col_idx] or '').strip()
        if not t:
            continue
        hdr_probe = row_cells[:5] if len(row_cells) >= 5 else row_cells
        if r < data_start and _excel_row_looks_like_header(hdr_probe):
            continue
        return t
    return ''


def _excel_cell_above(ws, row_1based, col_1based):
    """Cell text one row above (1-based Excel coordinates)."""
    if row_1based < 2 or not col_1based:
        return ''
    return _ws_cell(ws, row_1based - 1, col_1based).strip()


def _excel_vertical_merge_fill(ws, start_row_1based, col_1based, data_start_1based):
    """Last non-empty in ``col_1based`` scanning upward from ``start_row_1based``."""
    if not col_1based or start_row_1based < 1:
        return ''
    lo = max(1, int(data_start_1based))
    max_r = min(int(start_row_1based), ws.max_row or start_row_1based)
    for r in range(max_r, lo - 1, -1):
        t = _ws_cell(ws, r, col_1based).strip()
        if t:
            return t
    return ''


def _normalize_word_matrix_header_cell(raw):
    s = (raw or '').strip().lower()
    s = re.sub(r'\s+', ' ', s)
    s = _hdr_strip_leading_enum(s)
    s = re.sub(r'\s*\(\d+\)\s*$', '', s)
    return s.strip()


def _classify_word_matrix_header_cell(norm):
    """Bucket a single header cell for OPMM-style Word monitor tables."""
    if not norm:
        return 'OTHER'
    if norm in ('#', 'no.', 'no', 'item') or re.match(r'^no\.?\s*\d*$', norm):
        return 'ROW_NUM'
    if 'remark' in norm:
        return 'REMARKS'
    if 'variance' in norm:
        return 'VARIANCE'
    if 'accomplishment' in norm or ('actual' in norm and 'accomplish' in norm.replace(' ', '')):
        return 'ACCOMPLISHMENT'
    if 'indicator' in norm or ('performance' in norm and 'indicator' in norm):
        return 'INDICATOR'
    if 'strategy' in norm:
        return 'STRATEGY'
    if 'program' in norm and ('activity' in norm or 'project' in norm or 'pap' in norm):
        return 'PAP'
    if 'outcome' in norm:
        return 'OUTCOME'
    return 'OTHER'


def _word_monitor_row_is_likely_data(ind_text, ac_text):
    """True when this row looks like a KPI body line, not a column-title header row."""
    ind = (ind_text or '').strip()
    ac = (ac_text or '').strip()
    if not ind or not ac:
        return False
    # Strip leading enum like ``(5) Performance Indicator/s`` before matching header phrases:
    # BatStateU column headers carry ``(N)`` prefixes that the parens-target heuristic below
    # would otherwise mis-read as a KPI target and drop the real first KPI row.
    il = _hdr_strip_leading_enum(ind.lower())
    al = _hdr_strip_leading_enum(ac.lower())
    if 'performance indicator' in il and 'based on' in il:
        return False
    if re.match(r'^\s*(performance\s+)?indicator(?:/s|s)?\b', il) and re.match(
        r'^\s*(actual\s+)?accomplishment', al
    ):
        return False
    if il in ('indicator', 'indicators', 'performance indicator/s', 'performance indicators'):
        return False
    if al in ('actual accomplishment', 'actual accomplishments', 'accomplishment', 'accomplishments'):
        return False
    # After enum-strip, parenthesized digits are the only ``(N)`` left — those are real targets.
    ind_for_paren = _hdr_strip_leading_enum(ind)
    ac_for_paren = _hdr_strip_leading_enum(ac)
    if re.search(r'\(\s*\d+(?:\.\d+)?\s*\)', ind_for_paren) and re.search(
        r'\(\s*\d+(?:\.\d+)?\s*\)', ac_for_paren
    ):
        return True
    if re.search(r'\(\s*\d', ind_for_paren) and len(ac_for_paren) > 40:
        return True
    if len(ac_for_paren) > 120:
        return True
    return False


def _detect_word_matrix_column_map(table, max_scan_rows=12):
    """Infer 0-based column indices from header row text (handles leading # / shifted layouts).

    Returns ``(mapping, header_row_index)`` where ``mapping`` keys are
    OUTCOME, STRATEGY, PAP, INDICATOR, ACCOMPLISHMENT, VARIANCE, REMARKS.
    ``header_row_index`` is -1 when no header row was identified (legacy layout).
    """
    default_map = {
        'OUTCOME': 0,
        'STRATEGY': 1,
        'PAP': 2,
        'INDICATOR': 3,
        'ACCOMPLISHMENT': 4,
        'VARIANCE': 5,
        'REMARKS': 6,
    }
    logic_keys = frozenset(
        {'OUTCOME', 'STRATEGY', 'PAP', 'INDICATOR', 'ACCOMPLISHMENT', 'VARIANCE', 'REMARKS'}
    )
    best_score = -1
    best_map = dict(default_map)
    best_row = -1
    for ri in range(min(max_scan_rows, len(table.rows))):
        cells = _word_table_row_cells_grid(table, ri)
        if sum(1 for c in cells if (c or '').strip()) < 4:
            continue
        col_types = [_classify_word_matrix_header_cell(_normalize_word_matrix_header_cell(c)) for c in cells]
        present = {col_types[j] for j in range(len(col_types)) if col_types[j] in logic_keys}
        if 'INDICATOR' not in present:
            continue
        if 'ACCOMPLISHMENT' not in present and 'VARIANCE' not in present:
            continue
        # Body rows often omit the Outcome column (merged); title rows name it explicitly.
        if 'OUTCOME' not in present:
            continue
        mapping = {}
        for j, ct in enumerate(col_types):
            if ct in logic_keys and ct not in mapping:
                mapping[ct] = j
        ind_j = mapping.get('INDICATOR')
        ac_j = mapping.get('ACCOMPLISHMENT')
        ind_txt = cells[ind_j] if ind_j is not None and ind_j < len(cells) else ''
        ac_txt = cells[ac_j] if ac_j is not None and ac_j < len(cells) else ''
        if _word_monitor_row_is_likely_data(ind_txt, ac_txt):
            continue
        score = len(present) * 2
        if 'OUTCOME' in present:
            score += 2
        if 'STRATEGY' in present:
            score += 1
        if 'VARIANCE' in present:
            score += 2
        if 'REMARKS' in present:
            score += 1
        if score > best_score:
            best_score = score
            best_map = mapping
            best_row = ri
    if best_score < 0:
        return default_map, -1
    merged = dict(default_map)
    merged.update(best_map)
    return merged, best_row


def _word_monitor_mapped_cells(cells, col_map):
    """Build ``[out, strat, pap, ind, acmp, var, rem]`` using detected column indices."""

    def col(key, default_idx):
        idx = col_map.get(key, default_idx)
        if idx is None or idx < 0:
            return ''
        if idx >= len(cells):
            return ''
        return (cells[idx] or '').strip()

    return [
        col('OUTCOME', 0),
        col('STRATEGY', 1),
        col('PAP', 2),
        col('INDICATOR', 3),
        col('ACCOMPLISHMENT', 4),
        col('VARIANCE', 5),
        col('REMARKS', 6),
    ]


def _explicit_met_unmet_from_accomplishment_tail(accomplishment_text):
    """When variance was mis-mapped or empty, use a final line that is only MET/UNMET."""
    if not accomplishment_text or not str(accomplishment_text).strip():
        return None
    lines = [L.strip() for L in str(accomplishment_text).splitlines() if L.strip()]
    for L in reversed(lines[-5:]):
        if re.fullmatch(r'(?i)MET\.?', L):
            return 'MET'
        if re.fullmatch(r'(?i)UNMET\.?', L):
            return 'UNMET'
    return None


def _hdr_strip_leading_enum(h):
    """Strip ``(2)`` / ``2.`` style prefixes from normalized header labels."""
    s = (h or '').strip()
    s = re.sub(r'^\(\d+\)\s*', '', s)
    s = re.sub(r'^\d+[\.\)]\s*', '', s)
    return s.strip()


def _parse_explicit_met_unmet_cell(text):
    if not text:
        return None
    s = str(text).strip()
    # OPMM "Variance" cells are numeric: >= 0 means the office hit / exceeded target (MET),
    # anything below 0 means it fell short (UNMET). Plain ``0`` (no sign) must still be MET.
    m = re.fullmatch(r'\s*([+-]?)\s*(\d+(?:\.\d+)?)\s*%?\s*', s)
    if m:
        sign, num = m.group(1), m.group(2)
        try:
            val = float(num)
        except ValueError:
            val = None
        if val is not None:
            if sign == '-' and val > 0:
                return 'UNMET'
            return 'MET'
    if re.search(r'(?i)\bUNMET\b', s):
        return 'UNMET'
    if re.search(r'(?i)\bMET\b', s):
        return 'MET'
    if re.search(
        r'(?i)\b('
        r'fully\s+accomplished|fully\s+achieved|target\s+exceeded|'
        r'on\s*[-\s]?track|on\s+target|exceeded\s+target|surpassed|compliant'
        r')\b',
        s,
    ):
        return 'MET'
    # LPC / FYDP cells often use short affirmations instead of the word MET.
    if re.fullmatch(r'(?i)\s*(yes|y|ok|done|passed|complied|satisfactory|achieved)\s*\.?', s):
        return 'MET'
    if re.fullmatch(r'(?i)\s*(no|n|x|failed|not\s+complied)\s*\.?', s):
        return 'UNMET'
    if re.fullmatch(r'(?i)\s*([✓✔☑]|\+\s*|1)\s*', s):
        return 'MET'
    if re.fullmatch(r'(?i)\s*([✗☒]|-\s*|0)\s*', s):
        return 'UNMET'
    return None


def _extract_development_area_heading(text):
    """Return a normalized 'Development Area: ...' label if present, else None."""
    if not text:
        return None
    s = str(text).strip()
    if not s:
        return None
    m = re.search(r'(?i)\bdevelopment\s*area\s*[:\-]\s*(.+)$', s)
    if not m:
        return None
    tail = m.group(1).strip()
    tail = re.sub(r'\s+', ' ', tail)
    # Keep a reasonably-sized label; we only need it for bucketing.
    return f'Development Area: {tail[:220]}' if tail else 'Development Area'


def _extract_opmm_section_banner(text):
    """Detect OPMM block titles like ``SOCIAL RESPONSIBILITY: Engineering Pathways…``.

    Matrices often use ``PILLAR NAME: subtitle`` without the words ``Development Area``.
    """
    if not text:
        return None
    s = re.sub(r'\s+', ' ', str(text).strip())
    if not s:
        return None
    if len(s) > 600:
        s = s[:600]
    m = re.match(
        r'(?i)^(social\s+responsibility|research\s+and\s+innovation|academic\s+leadership|'
        r'internationalization|advancing\s+interdisciplinarity|sustainability)\s*:\s*(.+)$',
        s,
    )
    if m:
        return f'{m.group(1).strip()}: {m.group(2).strip()}'[:400]
    m3 = re.search(
        r'(?i)\b(social\s+responsibility|research\s+and\s+innovation|academic\s+leadership|'
        r'internationalization|advancing\s+interdisciplinarity|sustainability)\s*:\s*(.+)$',
        s,
    )
    if m3:
        return f'{m3.group(1).strip()}: {m3.group(2).strip()}'[:400]
    m2 = re.search(
        r'(?i)\bdevelopment\s+area\s*:\s*'
        r'(social\s+responsibility|research\s+and\s+innovation|academic\s+leadership|'
        r'internationalization|advancing\s+interdisciplinarity|sustainability)\s*'
        r':?\s*(.*)$',
        s,
    )
    if m2:
        tail = (m2.group(2) or '').strip()
        if tail:
            return f'{m2.group(1).strip()}: {tail}'[:400]
        return m2.group(1).strip()[:200]
    return None


# Six official OPMM pillars (display title + compact needles for canonicalization on ingest).
_OPMM_CANONICAL_PILLARS = (
    ('Sustainability', ('sustainability',)),
    ('Academic Leadership', ('academic leadership', 'academicleadership')),
    ('Research and Innovation', ('research and innovation', 'researchandinnovation')),
    ('Internationalization', ('internationalization', 'internationalisation')),
    ('Social Responsibility', ('social responsibility', 'socialresponsibility')),
    (
        'Advancing Interdisciplinarity',
        (
            'advancing interdisciplinarity',
            'advancinginterdisciplinarity',
            'interdisciplinarity',
        ),
    ),
)


def canonicalize_opmm_dev_area_label(raw):
    """Normalize banners/harvest text to ``<Canonical pillar>: subtitle`` for stable bucketing."""
    if not raw or not str(raw).strip():
        return (raw or '').strip()
    s = re.sub(r'(?i)^\s*development\s+area\s*:\s*', '', str(raw).strip()).strip()
    comp = re.sub(r'[^a-z0-9]', '', s.lower())
    for display, needles in _OPMM_CANONICAL_PILLARS:
        needle_cs = [re.sub(r'[^a-z0-9]', '', n.lower()) for n in needles]
        needle_cs.append(re.sub(r'[^a-z0-9]', '', display.lower()))
        if not any(n and n in comp for n in needle_cs):
            continue
        pat = r'\s+'.join(re.escape(w) for w in display.split())
        m = re.search(rf'(?i){pat}\s*[:-]\s*(.+)$', s)
        if m and m.group(1).strip():
            return f'{display}: {m.group(1).strip()}'[:500]
        return display
    return s


def _scan_row_for_development_area(cells):
    """Scan a row for ``Development Area:`` text or an OPMM pillar banner."""
    if not cells:
        return None
    for c in cells:
        v = _extract_development_area_heading(c)
        if v:
            return v
        b = _extract_opmm_section_banner(_cell_value_str(c))
        if b:
            return b
    joined = ' '.join(_cell_value_str(c) for c in cells if c)
    return _extract_opmm_section_banner(joined)


def _compact_label_key(text):
    return re.sub(r'[^a-z0-9]', '', (text or '').lower())


def _accomplishment_is_non_kpi_placeholder(text):
    """Accomplishment cell empty, placeholder, or another pillar subtitle — not a KPI narrative."""
    t = (text or '').strip()
    if not t:
        return True
    if re.fullmatch(r'(?i)([-—–]|n\s*/\s*a|na|none|nil)\.?', t):
        return True
    if _extract_opmm_section_banner(t):
        return True
    return False


def _looks_like_performance_indicator(text):
    """True when ``text`` is a quantifiable KPI line, not an OPMM pillar/subtitle banner."""
    t = (text or '').strip()
    if not t or len(t) < 5:
        return False
    if not _extract_opmm_section_banner(t):
        return True
    if re.search(r'\(\s*\d+', t):
        return True
    if re.search(
        r'(?i)\b(number of|no\.?\s*of|percentage|percent|rate|ratio|conducted|submitted|'
        r'completed|achieved|participants?|programs?|sessions?|publications?|patents?|'
        r'facilitated|conducted|developed|established)\b',
        t,
    ):
        return True
    return False


def _row_has_substantive_kpi_payload(
    indicator_text, accomplishment_text, variance_text='', remarks_text=''
):
    """True when the row carries a real KPI (not only a repeated pillar title row)."""
    ind = (indicator_text or '').strip()
    ac = (accomplishment_text or '').strip()
    if _looks_like_performance_indicator(ind):
        return True
    if ac and not _accomplishment_is_non_kpi_placeholder(ac):
        if len(ac) >= 10 or extract_actual_for_monitor(ac) is not None:
            return True
    rem = (remarks_text or '').strip()
    if rem and len(rem) >= 16 and not _extract_opmm_section_banner(rem):
        return True
    var = (variance_text or '').strip()
    if var and re.match(r'^\s*[+-]\s*\d+\s*$', var):
        return True
    return False


def _indicator_cell_is_pillar_banner_not_kpi(indicator_text):
    """Indicator column holds a merged pillar title, not a quantifiable KPI line."""
    t = (indicator_text or '').strip()
    if not t:
        return False
    return bool(_extract_opmm_section_banner(t)) and not _looks_like_performance_indicator(t)


def _word_row_pick_indicator_text(cells, col_map):
    """Best performance-indicator string in a grid row (mapped column first, then scan)."""
    ind_i = col_map.get('INDICATOR', 3)
    if ind_i is not None and 0 <= ind_i < len(cells):
        t = (cells[ind_i] or '').strip()
        if _looks_like_performance_indicator(t):
            return t
    for c in cells:
        t = (c or '').strip()
        if _looks_like_performance_indicator(t):
            return t
    if ind_i is not None and 0 <= ind_i < len(cells):
        return (cells[ind_i] or '').strip()
    return ''


def _row_is_pure_opmm_section_banner_row(
    outcome_text,
    strategy_text,
    pap_text,
    indicator_text,
    accomplishment_text,
    variance_text='',
    remarks_text='',
):
    """True when the row is only a development-area / pillar title, not a KPI matrix line.

    Word exports often repeat ``SOCIAL RESPONSIBILITY: …`` across Outcome–Indicator columns
    with ``MET`` in Variance; that must not become a ``PerformanceRecord``.
    """
    if _row_has_substantive_kpi_payload(
        indicator_text, accomplishment_text, variance_text, remarks_text
    ):
        return False

    cells_hier = [
        (outcome_text or '').strip(),
        (strategy_text or '').strip(),
        (pap_text or '').strip(),
        (indicator_text or '').strip(),
    ]
    ind = (indicator_text or '').strip()
    ac = (accomplishment_text or '').strip()

    banners = [_extract_opmm_section_banner(c) for c in cells_hier if c]
    banners = [b for b in banners if b]
    if len(banners) >= 2 and len({_compact_label_key(b) for b in banners}) == 1:
        return True

    non_empty = [c for c in cells_hier if len(c) > 24]
    if len(non_empty) >= 2 and len({_compact_label_key(c) for c in non_empty}) == 1:
        return True

    if ind and not _looks_like_performance_indicator(ind):
        if _accomplishment_is_non_kpi_placeholder(ac):
            return True
        if _compact_label_key(ac) == _compact_label_key(ind):
            return True

    rem = (remarks_text or '').strip()
    var = (variance_text or '').strip()
    if ind and not _looks_like_performance_indicator(ind):
        if not ac and not rem and var in ('MET', 'UNMET'):
            return True
    return False


def record_is_pillar_banner_only(record):
    """Viewer/export guard: exclude ingested pillar title rows mistaken for KPIs."""
    ind = (record.indicator.description or '').strip()
    if not ind:
        return False
    if _looks_like_performance_indicator(ind):
        return False
    raw = (record.raw_actual_text or '').strip()
    var = (record.variance_text or '').strip()
    if not _extract_opmm_section_banner(ind):
        return False
    if not _accomplishment_is_non_kpi_placeholder(raw):
        return False
    if raw and _looks_like_performance_indicator(raw):
        return False
    if var in ('MET', 'UNMET') and _accomplishment_is_non_kpi_placeholder(raw):
        return True
    return _accomplishment_is_non_kpi_placeholder(raw)


def _should_skip_row_as_dev_area_banner_only(
    scan_cells,
    indicator_text,
    accomplishment_text,
    variance_text='',
    remarks_text='',
    *,
    outcome_text='',
    strategy_text='',
    pap_text='',
):
    """When a row matches a dev-area / pillar banner, skip only if it is not a matrix KPI row.

    OPMM tables often put ``RESEARCH AND INNOVATION: …`` in the Outcome cell on every KPI line.
    ``_scan_row_for_development_area`` then matches the same row; treating it as a banner-only
    row and ``continue`` dropped all such indicators (e.g. 4 in the file → 0 or 2 ingested).

    Word vertical merges often leave the **indicator** cell empty on continuation rows while
    accomplishment / variance still hold the KPI — those rows must not be classified as
    banner-only (the old logic returned True whenever ``indicator`` was empty).
    """
    c0 = outcome_text or ((scan_cells[0] if scan_cells else '') or '')
    c1 = strategy_text or ((scan_cells[1] if len(scan_cells) > 1 else '') or '')
    c2 = pap_text or ((scan_cells[2] if len(scan_cells) > 2 else '') or '')
    if _row_is_pure_opmm_section_banner_row(
        c0, c1, c2, indicator_text, accomplishment_text, variance_text, remarks_text
    ):
        return True
    if _word_row_has_kpi_signal(indicator_text, accomplishment_text, variance_text, remarks_text):
        return False
    ind = (indicator_text or '').strip()
    ac = (accomplishment_text or '').strip()
    var = (variance_text or '').strip()
    rem = (remarks_text or '').strip()
    if ind or ac or var or rem:
        return False
    return True


def _should_skip_row_as_dev_area_banner_only_lpc(
    scan_cells, indicator_text, annual_text, acmp_text, status_text=''
):
    """LPC layout: skip only pure banner rows with no indicator / target / accomplishment / status."""
    if _word_row_has_kpi_signal(indicator_text, acmp_text, status_text, ''):
        return False
    ind = (indicator_text or '').strip()
    ann = (annual_text or '').strip()
    ac = (acmp_text or '').strip()
    st = (status_text or '').strip()
    if ind or ann or ac or st:
        if ind and (ann or ac):
            probe = list(scan_cells[:5]) if len(scan_cells) >= 5 else list(scan_cells)
            if not _excel_row_looks_like_header(probe):
                return False
        return False
    return True


def _harvest_lpc_development_area(ws, header_row, max_cols=24):
    """Title text above the table: ``DEVELOPMENT AREA: …`` or ``SOCIAL RESPONSIBILITY: …``."""
    mc = max(1, min(max_cols, ws.max_column or max_cols))
    for r in range(1, max(1, header_row)):
        parts = []
        for c in range(1, mc + 1):
            parts.append(_cell_value_str(ws.cell(row=r, column=c).value))
        line = ' '.join(t for t in parts if t)
        m = re.search(r'(?i)development\s+area\s*:?\s*(.+)', line)
        if m:
            return canonicalize_opmm_dev_area_label(m.group(1).strip()[:800])
        banner = _extract_opmm_section_banner(line)
        if banner:
            return canonicalize_opmm_dev_area_label(banner[:800])
    return ''


def _map_lpc_wide_columns(hdrs):
    """Map normalized header strings (per column, 0-based index) to 1-based column indices."""
    n = len(hdrs)

    def pick(pred):
        for i in range(n):
            if pred(hdrs[i]):
                return i + 1
        return None

    annual = pick(
        lambda h: ('annual' in h and 'quantifiable' in h)
        or ('quantifiable' in h and 'target' in h and 'quarter' not in h)
    )
    acmp = pick(
        lambda h: 'accomplishment' in h
        and ('date' in h or 'sum' in h or 'quarterly' in h)
    ) or pick(
        lambda h: 'actual' in h and 'accomplish' in h
    ) or pick(
        lambda h: ('accomplishment' in h or 'accomplishments' in h)
        and 'annual' not in h
        and 'quantifiable' not in h
        and 'status' not in h
    )
    status_c = pick(
        lambda h: 'status' in h
        and ('met' in h or 'unmet' in h or 'accomplishment' in h)
    ) or pick(lambda h: h == 'variance' or h.startswith('variance'))
    outcome = pick(
        lambda h: (h == 'outcome' or h.startswith('outcome'))
        and 'development' not in h
    )
    strategy = pick(lambda h: 'strategy' in h and 'based' in h) or pick(
        lambda h: h.startswith('strategy')
    )
    program = pick(
        lambda h: 'program' in h and ('activity' in h or 'project' in h)
    )
    subpap = pick(
        lambda h: ('sub' in h and 'pap' in h) and 'program' not in h
    )
    ind_sub = pick(
        lambda h: 'indicator' in h and 'submitted' in h and 'operational' in h
    )
    ind_cas = pick(
        lambda h: 'indicator' in h and ('cascaded' in h or 'template' in h)
    )
    indicator = ind_sub or ind_cas or pick(
        lambda h: 'performance' in h and 'indicator' in h
    )
    office_c = pick(lambda h: 'concerned' in h and 'office' in h)
    return {
        'outcome': outcome,
        'strategy': strategy,
        'program': program,
        'sub_pap': subpap,
        'indicator': indicator,
        'annual_target': annual,
        'accomplishment': acmp,
        'status': status_c,
        'office': office_c,
    }


def _score_lpc_wide_header(hdrs):
    cols = _map_lpc_wide_columns(hdrs)
    score = 0
    if cols['annual_target']:
        score += 4
    if cols['accomplishment']:
        score += 4
    if cols['indicator']:
        score += 2
    if cols['outcome'] or cols['strategy']:
        score += 1
    if cols['status']:
        score += 1
    if cols['program']:
        score += 1
    return score, cols


def _detect_lpc_wide_layout(ws, scan_rows=45, max_cols=24):
    """LPC / FYDP wide workbook: annual target + accomplishment + indicator columns."""
    mc = max(max_cols, ws.max_column or 0, 16)
    last_r = min(scan_rows, ws.max_row or scan_rows)
    best_score, best_row, best_cols = 0, None, None
    for r in range(1, last_r + 1):
        hdrs = []
        for c in range(1, mc + 1):
            hdrs.append(_normalize_excel_header_text(ws.cell(row=r, column=c).value))
        score, cols = _score_lpc_wide_header(hdrs)
        if score > best_score:
            best_score, best_row, best_cols = score, r, cols
    if best_score < 6 or not best_cols:
        return None
    if not (best_cols.get('annual_target') and best_cols.get('accomplishment')):
        return None
    if not best_cols.get('indicator'):
        return None
    return {'header_row': best_row, 'cols': best_cols, 'max_col': mc}


def _map_opmm_matrix_columns(hdrs):
    """Operational Plan Monitoring Matrix: Outcome, Strategy, PAP, Indicator, Actual, Variance, Remarks.

    Typical of BatStateU OVCRDES Q4 matrices without a separate *Annual Quantifiable Target* column.
    """
    n = len(hdrs)
    ch = [_hdr_strip_leading_enum(h) for h in hdrs]

    def pick(pred):
        for i in range(n):
            if pred(ch[i]):
                return i + 1
        return None

    outcome = pick(
        lambda h: ('outcome' in h or h.startswith('outcome'))
        and 'development' not in h
        and 'operational' not in h
    )
    strategy = pick(lambda h: 'strategy' in h)
    program = pick(
        lambda h: ('pap' in h or 'program' in h)
        and (
            'activity' in h
            or 'project' in h
            or 'action' in h
            or 'step' in h
            or 'pap' in h
        )
    )
    subpap = pick(lambda h: ('sub' in h and 'pap' in h) and 'program' not in h)
    indicator = pick(lambda h: 'performance' in h and 'indicator' in h) or pick(
        lambda h: bool(re.search(r'\bindicator', h))
    )
    acmp = pick(lambda h: 'actual' in h and 'accomplish' in h) or pick(
        lambda h: 'accomplishment' in h and 'annual' not in h and 'quantifiable' not in h
    )
    status = pick(lambda h: h == 'variance' or h.startswith('variance')) or pick(
        lambda h: 'status' in h
        and ('met' in h or 'unmet' in h or 'accomplish' in h or 'variance' in h)
    )
    remarks = pick(lambda h: 'remark' in h)
    office_c = pick(lambda h: 'concerned' in h and 'office' in h)
    return {
        'outcome': outcome,
        'strategy': strategy,
        'program': program,
        'sub_pap': subpap,
        'indicator': indicator,
        'accomplishment': acmp,
        'status': status,
        'remarks': remarks,
        'office': office_c,
    }


def _score_opmm_matrix_header(hdrs):
    cols = _map_opmm_matrix_columns(hdrs)
    score = 0
    hier = 0
    if cols['indicator']:
        score += 4
    if cols['accomplishment']:
        score += 4
    if cols['outcome']:
        hier += 1
    if cols['strategy']:
        hier += 1
    if cols['program']:
        hier += 1
    if hier:
        score += min(hier, 3)
    if cols['status']:
        score += 2
    return score, cols


def _detect_opmm_matrix_layout(ws, scan_rows=50, max_cols=28):
    """Detect OPMM-style monitoring tables (Actual Accomplishments + Variance, no annual target column)."""
    mc = max(max_cols, ws.max_column or 0, 12)
    last_r = min(scan_rows, ws.max_row or scan_rows)
    best_score, best_row, best_cols = 0, None, None
    for r in range(1, last_r + 1):
        hdrs = []
        for c in range(1, mc + 1):
            hdrs.append(_normalize_excel_header_text(ws.cell(row=r, column=c).value))
        score, cols = _score_opmm_matrix_header(hdrs)
        if score > best_score:
            best_score, best_row, best_cols = score, r, cols
    if best_score < 9 or not best_cols:
        return None
    if not (best_cols.get('indicator') and best_cols.get('accomplishment')):
        return None
    if not (best_cols.get('outcome') or best_cols.get('strategy') or best_cols.get('program')):
        return None
    return {'header_row': best_row, 'cols': best_cols, 'max_col': mc}


def _ingest_monitor_row_extended(
    office_obj,
    outcome_name,
    strategy_name,
    pap_name,
    indicator_name,
    target_source_text,
    accomplishment_text,
    explicit_status,
    current_q,
    current_y,
    variance_text='',
):
    """Persist hierarchy + KPI; target parsed from ``target_source_text`` (not only indicator)."""
    if not (indicator_name or '').strip():
        return False
    if not (outcome_name or strategy_name or pap_name):
        return False

    out, _ = StrategicLevel.objects.get_or_create(
        name=outcome_name or 'Outcome',
        level_type='OUTCOME',
    )
    st, _ = StrategicLevel.objects.get_or_create(
        name=strategy_name or 'Strategy',
        level_type='STRATEGY',
        parent=out,
    )
    pap, _ = StrategicLevel.objects.get_or_create(
        name=pap_name or 'Program / PAP',
        level_type='PAP',
        parent=st,
        office=office_obj,
    )

    indicator, _ = Indicator.objects.get_or_create(pap=pap, description=indicator_name or 'Indicator')

    target_val = extract_target_for_monitor(target_source_text or '')
    if target_val is None:
        target_val = extract_target_for_monitor(indicator_name or '')
    actual_val = extract_actual_for_monitor(accomplishment_text or '')

    v_raw = (variance_text or '').strip()
    exp = explicit_status if explicit_status in ('MET', 'UNMET') else None
    if exp is None and v_raw:
        exp = _parse_explicit_met_unmet_cell(v_raw)
        if exp not in ('MET', 'UNMET'):
            exp = None
    if exp is None:
        exp = _explicit_met_unmet_from_accomplishment_tail(accomplishment_text or '')
        if exp not in ('MET', 'UNMET'):
            exp = None
    defaults = {
        'target_value': target_val,
        'actual_value': actual_val,
        'raw_actual_text': accomplishment_text or '',
        'variance_text': v_raw,
        'explicit_status': exp,
    }

    PerformanceRecord.objects.update_or_create(
        indicator=indicator,
        quarter=current_q,
        year=current_y,
        defaults=defaults,
    )
    return True


def _ingest_monitor_row(office_obj, cells, current_q, current_y):
    """Narrow template: outcome…indicator, accomplishment; optional variance + remarks columns."""
    if len(cells) < 5:
        return False
    var_txt = (cells[5] or '').strip() if len(cells) > 5 else ''
    rem_txt = (cells[6] or '').strip() if len(cells) > 6 else ''
    acmp = (cells[4] or '').strip()
    if rem_txt:
        acmp = f'{acmp}\n\n— Remarks —\n{rem_txt}' if acmp else rem_txt
    exp = _parse_explicit_met_unmet_cell(var_txt) if var_txt else None
    return _ingest_monitor_row_extended(
        office_obj,
        cells[0],
        cells[1],
        cells[2],
        cells[3],
        cells[3],
        acmp,
        exp,
        current_q,
        current_y,
        variance_text=var_txt,
    )


def _word_row_has_kpi_signal(indicator_text, accomplishment_text, variance_text, remarks_text):
    """Heuristic: row likely carries KPI content even if hierarchy columns are blank."""
    ind = (indicator_text or '').strip()
    ac = (accomplishment_text or '').strip()
    var = (variance_text or '').strip()
    rem = (remarks_text or '').strip()
    if ind and len(ind) >= 5:
        if _looks_like_performance_indicator(ind):
            return True
    if ac:
        if len(ac) >= 8:
            return True
        if extract_actual_for_monitor(ac) is not None:
            return True
        if _parse_explicit_met_unmet_cell(ac) in ('MET', 'UNMET'):
            return True
        if re.search(r'(?i)\b(met|unmet|n\s*/\s*a|n/?a)\b', ac):
            return True
    if var and (_parse_explicit_met_unmet_cell(var) in ('MET', 'UNMET') or re.match(r'^\s*[+-]\s*\d+\s*$', var)):
        if ind and not _looks_like_performance_indicator(ind) and _accomplishment_is_non_kpi_placeholder(ac):
            return False
        return True
    if rem and len(rem) >= 8:
        return True
    return False


def _ingest_word_monitor_table(table, office_obj, current_q, current_y):
    """Ingest one Word table and return how many performance rows were written."""
    if not table or not getattr(table, 'rows', None):
        return 0, {}
    rows_written = 0
    skipped = {'no_signal': 0, 'no_context': 0, 'no_indicator': 0}
    col_map, header_row_idx = _detect_word_matrix_column_map(table)
    # When a real matrix header row is detected, the first data row is immediately below it.
    # Do not force ``data_start >= 3`` here: OVCRDES-style tables often place the header on
    # row 0 or 1 with KPIs starting at index 1 or 2; ``max(3, header + 1)`` skipped those rows
    # so only the later indicators were ingested (e.g. 4 in file → 2 in DB).
    if header_row_idx >= 0:
        data_start = header_row_idx + 1
    else:
        data_start = 3  # legacy: no detected header — assume ~2 title rows before matrix body

    dev_area_current = None
    last_out, last_strat, last_prog = '', '', ''
    last_ind_txt = ''
    seen_matrix_indicator_keys = set()
    for i, row in enumerate(table.rows):
        cells = _word_table_row_cells_grid(table, i)
        if sum(1 for c in cells if (c or '').strip()) < 2:
            continue
        if i < data_start:
            maybe_dev = _scan_row_for_development_area(cells)
            if maybe_dev:
                dev_area_current = canonicalize_opmm_dev_area_label(maybe_dev)
            continue

        std = _word_monitor_mapped_cells(cells, col_map)
        raw0, raw1, raw2 = std[0], std[1], std[2]
        c3 = std[3]
        c4, c5, c6 = std[4], std[5], std[6]
        oi = col_map.get('OUTCOME', 0)
        si = col_map.get('STRATEGY', 1)
        pi = col_map.get('PAP', 2)
        # Word vertical merges often leave continuation rows blank; without fill we keep the
        # previous row's Strategy/PAP and attach the next KPI to the wrong PAP block.
        c0 = (raw0 or '').strip()
        c1 = (raw1 or '').strip()
        c2 = (raw2 or '').strip()
        c3_raw = (c3 or '').strip()
        # Merged pillar rows clone the banner into the indicator column; keep accomplishment.
        if _indicator_cell_is_pillar_banner_not_kpi(c3_raw) and _row_has_substantive_kpi_payload(
            '', c4, c5, c6
        ):
            c3 = ''
        if not _looks_like_performance_indicator((c3 or '').strip()):
            picked = _word_row_pick_indicator_text(cells, col_map)
            if picked and not _indicator_cell_is_pillar_banner_not_kpi(picked):
                c3 = picked
        if i >= data_start:
            if not c0:
                c0 = _word_table_vertical_merge_fill(table, i - 1, oi, data_start, header_row_idx)
            if not c1:
                c1 = _word_table_vertical_merge_fill(table, i - 1, si, data_start, header_row_idx)
            if not c1 and i > data_start:
                # Never read the column-header row via ``_word_cell_above`` on the *first* body row:
                # the cell above ``data_start`` is often the literal ``Strategy`` / ``Outcome`` label,
                # which produced phantom KPI rows and corrupted ``PerformanceRecord`` upserts.
                c1 = _word_cell_above(table, i, si)
            # OPMM Word exports often store merged PAP text only on the first row of a span;
            # the row directly above may be empty, so ``_word_cell_above`` is not enough.
            # The old strategy-change gate also set ``c2 = ''`` and cleared ``last_prog`` on
            # false-positive strategy diffs — dropping the first KPI rows under a PAP block.
            if not (raw2 or '').strip():
                c2 = _word_table_vertical_merge_fill(table, i - 1, pi, data_start, header_row_idx)

        # Pillar / subtitle row immediately above the detected column-header row: some BatStateU
        # templates vertically merge that banner into the first KPI row, so python-docx leaves
        # Outcome/Strategy/PAP empty on the first body row while the visible anchor sits one row
        # above the ``Outcome (2)`` header line.
        if (
            i == data_start
            and header_row_idx >= 1
            and not (raw0 or '').strip()
            and not (raw1 or '').strip()
            and not (raw2 or '').strip()
        ):
            prow = _word_table_row_cells_grid(table, header_row_idx - 1)
            hdr_probe = prow[:5] if len(prow) >= 5 else prow
            if len(prow) >= 4 and not _excel_row_looks_like_header(hdr_probe):
                joined = ' '.join(x for x in prow if x)
                if _scan_row_for_development_area(prow) or _extract_opmm_section_banner(joined):
                    pstd = _word_monitor_mapped_cells(prow, col_map)
                    if not c0 and (pstd[0] or '').strip():
                        c0 = (pstd[0] or '').strip()
                    if not c1 and (pstd[1] or '').strip():
                        c1 = (pstd[1] or '').strip()
                    if not c2 and (pstd[2] or '').strip():
                        c2 = (pstd[2] or '').strip()

        scan_for_header = cells[: min(8, len(cells))]
        if _excel_row_looks_like_header(scan_for_header[:5] if len(scan_for_header) >= 5 else scan_for_header):
            continue

        indicator_hdr = (c3 or '').strip().lower()
        accom_hdr = (c4 or '').strip().lower()
        if indicator_hdr in (
            'performance indicator/s',
            'performance indicators',
            'indicator',
            'indicators',
        ):
            continue
        if accom_hdr in ('actual accomplishment', 'actual accomplishments', 'accomplishment'):
            continue
        scan_cells = [c0, c1, c2, c3, c4, c5, c6]
        maybe_dev = _scan_row_for_development_area(scan_cells)
        if maybe_dev:
            dev_area_current = canonicalize_opmm_dev_area_label(maybe_dev)
            if _should_skip_row_as_dev_area_banner_only(
                scan_cells,
                c3,
                c4,
                c5,
                c6,
                outcome_text=c0,
                strategy_text=c1,
                pap_text=c2,
            ):
                last_out, last_strat, last_prog = '', '', ''
                last_ind_txt = ''
                continue

        has_parent_cols = bool(c0 or c1 or c2 or c3)
        if not has_parent_cols and not (c4 or c5 or c6) and not (last_out or last_strat or last_prog):
            skipped['no_signal'] += 1
            continue
        if not has_parent_cols and not (last_out or last_strat or last_prog):
            if _word_row_has_kpi_signal(c3, c4, c5, c6):
                # Salvage KPI-like orphan rows by attaching defaults so they aren't silently dropped.
                last_out, last_strat, last_prog = 'Outcome', 'Strategy', 'Program / PAP'
            else:
                skipped['no_context'] += 1
                continue

        if c0 and not _indicator_cell_is_pillar_banner_not_kpi(c0):
            if _matrix_label_changed(last_out, c0):
                last_strat, last_prog = '', ''
                last_ind_txt = ''
            last_out = c0
        if c1 and not _indicator_cell_is_pillar_banner_not_kpi(c1):
            if _matrix_label_changed(last_strat, c1):
                last_prog = ''
                last_ind_txt = ''
            last_strat = c1
        if c2 and not _indicator_cell_is_pillar_banner_not_kpi(c2):
            if _matrix_label_changed(last_prog, c2):
                last_ind_txt = ''
            last_prog = c2

        row_cells = [last_out, last_strat, last_prog, '', c4, c5, c6]
        if c3:
            last_ind_txt = c3
            row_cells[3] = c3
        elif (c4 or c5 or c6) and (
            (last_out or last_strat or last_prog)
            or _row_has_substantive_kpi_payload('', c4, c5, c6)
        ):
            first_line = ''
            if c4:
                first_line = c4.strip().splitlines()[0].strip()
            if not first_line and c6:
                first_line = c6.strip().splitlines()[0].strip()
            tail = first_line
            if last_ind_txt:
                # Always tag physical row: two continuations can share the same 220-char prefix
                # (or identical first lines), which would otherwise yield one ``Indicator`` row.
                row_cells[3] = (
                    f'{last_ind_txt} — {tail[:220]} ·r{i + 1}'
                    if tail
                    else f'{last_ind_txt} — (continuation, row {i + 1})'
                )
            else:
                base = tail if tail else 'Performance indicator'
                row_cells[3] = f'{base} (table row {i + 1})'
        else:
            row_cells[3] = c3

        if _indicator_cell_is_pillar_banner_not_kpi(row_cells[3]):
            row_cells[3] = ''
        for hi in (0, 1, 2):
            if _indicator_cell_is_pillar_banner_not_kpi(row_cells[hi]):
                row_cells[hi] = ''
        if not (row_cells[0] or row_cells[1] or row_cells[2]) and _row_has_substantive_kpi_payload(
            row_cells[3], c4, c5, c6
        ):
            row_cells[0], row_cells[1], row_cells[2] = (
                last_out or 'Outcome',
                last_strat or 'Strategy',
                last_prog or 'Program / PAP',
            )

        if dev_area_current:
            out_txt = (row_cells[0] or '').strip()
            d_c = re.sub(r'[^a-z0-9]', '', dev_area_current.lower())
            o_c = re.sub(r'[^a-z0-9]', '', out_txt.lower())
            if out_txt and d_c[:24] not in o_c:
                row_cells[0] = f'{dev_area_current} — {out_txt or "Outcome"}'

        # One PerformanceRecord per (indicator, quarter, year). Word merges often duplicate the
        # same indicator text on the next row while accomplishment differs; without a distinct
        # ``Indicator.description`` the second row overwrites the first (4 matrix lines → 3 rows).
        hier_key = (
            (row_cells[0] or '').strip(),
            (row_cells[1] or '').strip(),
            (row_cells[2] or '').strip(),
        )
        dedup_desc = (row_cells[3] or '').strip()
        dedup_key = (hier_key, dedup_desc)
        suf = 1
        while dedup_key in seen_matrix_indicator_keys:
            suf += 1
            alt = f'{dedup_desc} · row {i + 1}' if suf == 2 else f'{dedup_desc} · row {i + 1} ({suf})'
            dedup_key = (hier_key, alt)
        if suf > 1:
            row_cells[3] = dedup_key[1]
        seen_matrix_indicator_keys.add(dedup_key)

        if not (row_cells[3] or '').strip():
            picked = _word_row_pick_indicator_text(cells, col_map)
            if picked and not _indicator_cell_is_pillar_banner_not_kpi(picked):
                row_cells[3] = picked
            elif (c4 or '').strip() and not _accomplishment_is_non_kpi_placeholder(c4):
                row_cells[3] = (c4 or '').strip().splitlines()[0].strip()[:240]
            elif (c6 or '').strip():
                row_cells[3] = (c6 or '').strip().splitlines()[0].strip()[:240]

        if not (row_cells[3] or '').strip():
            skipped['no_indicator'] += 1
            continue

        while len(row_cells) < 7:
            row_cells.append('')
        if _ingest_monitor_row(office_obj, row_cells[:7], current_q, current_y):
            rows_written += 1
    return rows_written, skipped


def _split_word_paragraph_table_row(text):
    """Parse table-like paragraph rows delimited by ``|`` or tabs; returns 7 cells or ``None``."""
    if not text or not str(text).strip():
        return None
    s = str(text).strip()
    parts = None
    if '|' in s:
        cand = [p.strip() for p in s.split('|')]
        if len([p for p in cand if p]) >= 5:
            parts = cand
    elif '\t' in s:
        cand = [p.strip() for p in s.split('\t')]
        if len([p for p in cand if p]) >= 5:
            parts = cand
    if not parts:
        return None
    while len(parts) < 7:
        parts.append('')
    return parts[:7]


def _ingest_word_monitor_paragraph_rows(doc, office_obj, current_q, current_y):
    """Fallback ingest for table-like KPI rows that appear as paragraphs, not true tables."""
    rows_written = 0
    last_out, last_strat, last_prog, last_ind = '', '', '', ''
    for p in doc.paragraphs:
        row = _split_word_paragraph_table_row((p.text or '').strip())
        if not row:
            continue
        c0, c1, c2, c3, c4, c5, c6 = [x.strip() for x in row[:7]]
        maybe_dev = _scan_row_for_development_area(row)
        if maybe_dev:
            if _should_skip_row_as_dev_area_banner_only(
                row,
                c3,
                c4,
                c5,
                c6,
                outcome_text=c0,
                strategy_text=c1,
                pap_text=c2,
            ):
                last_out, last_strat, last_prog, last_ind = '', '', '', ''
                continue
        if not any([c0, c1, c2, c3, c4, c5, c6]):
            continue
        if _excel_row_looks_like_header([c0, c1, c2, c3, c4]):
            continue
        if c0:
            if _matrix_label_changed(last_out, c0):
                last_strat, last_prog, last_ind = '', '', ''
            last_out = c0
        if c1:
            if _matrix_label_changed(last_strat, c1):
                last_prog, last_ind = '', ''
            last_strat = c1
        if c2:
            if _matrix_label_changed(last_prog, c2):
                last_ind = ''
            last_prog = c2
        if c3:
            last_ind = c3
        elif (c4 or c5 or c6) and last_ind:
            c3 = f'{last_ind} — (paragraph row)'

        if not c3:
            continue
        _ingest_monitor_row(
            office_obj,
            [last_out, last_strat, last_prog, c3, c4, c5, c6],
            current_q,
            current_y,
        )
        rows_written += 1
    return rows_written


def ingest_word_monitor_with_stats(file_path, office_name="OVCRDES", quarter=None, year=None, extra_hint=''):
    """Parse a Word monitor table into performance rows.

    ``office_name`` ties the data to the user's office. ``quarter`` / ``year``
    (when supplied) override values detected from the document title text.
    ``extra_hint`` (e.g. original filename) is prepended for period detection.

    Returns ``{'rows_written': int, 'breakdown': [(label, count), ...]}``.
    """
    doc = docx.Document(file_path)
    header_text = _harvest_period_hint_text(doc)
    if extra_hint:
        header_text = f'{extra_hint.strip()}\n{header_text}'
    auto_q, auto_y = _detect_quarter_year(header_text)
    current_q = quarter if quarter is not None else (auto_q or 1)
    current_y = year if year is not None else (auto_y or 2025)

    office_obj, _ = Office.objects.get_or_create(name=office_name)

    rows_written = 0
    breakdown = []
    skipped_total = {'no_signal': 0, 'no_context': 0, 'no_indicator': 0}
    for table in doc.tables:
        idx = len(breakdown) + 1
        c, sk = _ingest_word_monitor_table(table, office_obj, current_q, current_y)
        rows_written += c
        breakdown.append((f'Table {idx}', c))
        skipped_total['no_signal'] += sk.get('no_signal', 0)
        skipped_total['no_context'] += sk.get('no_context', 0)
        skipped_total['no_indicator'] += sk.get('no_indicator', 0)
    para_rows = _ingest_word_monitor_paragraph_rows(doc, office_obj, current_q, current_y)
    rows_written += para_rows
    breakdown.append(('Paragraph rows', para_rows))
    breakdown.append(('Skipped (no signal)', skipped_total['no_signal']))
    breakdown.append(('Skipped (no context)', skipped_total['no_context']))
    breakdown.append(('Skipped (no indicator)', skipped_total.get('no_indicator', 0)))
    return {'rows_written': rows_written, 'breakdown': breakdown}


def ingest_word_monitor(file_path, office_name="OVCRDES", quarter=None, year=None, extra_hint=''):
    """Backward-compatible wrapper returning only the row count."""
    return ingest_word_monitor_with_stats(
        file_path,
        office_name=office_name,
        quarter=quarter,
        year=year,
        extra_hint=extra_hint,
    )['rows_written']


def _excel_row_looks_like_header(cells):
    """True when row looks like column titles (whole-cell labels), not data."""
    if not cells or len(cells) < 3:
        return False
    canon = [(c or '').strip().lower() for c in cells[:5]]
    if not any(canon):
        return False
    if canon[0] in ('outcome', 'outcomes', '#', 'no.', 'no', 'item'):
        return True
    labels = {
        'outcome',
        'outcomes',
        'strategy',
        'pap',
        'indicator',
        'indicators',
        'accomplishment',
        'accomplishments',
        'performance',
    }
    # Match whole cell text only — narrative rows contain words like "strategy"
    # inside longer phrases and must not be treated as headers.
    return sum(1 for c in canon if c in labels) >= 3


def _excel_effective_max_row(ws, data_start_hint=1):
    """Some exports report a too-small ``max_row``; widen scan so data is not skipped."""
    mr = ws.max_row or 0
    end = max(mr, data_start_hint)
    if mr < data_start_hint + 3 or end <= data_start_hint:
        end = max(end, data_start_hint + 2000)
    return end


def _excel_data_start_row(ws, max_scan=260):
    """Find first data row: A–E look like content, not merged title/header labels."""
    scan_cap = max(max_scan, ws.max_row or 0, 260)
    max_r = min(_excel_effective_max_row(ws, 1), scan_cap)
    for r in range(1, max_r + 1):
        cells = [
            _cell_value_str(ws.cell(row=r, column=c).value) for c in range(1, 6)
        ]
        if _excel_row_looks_like_header(cells):
            continue
        ind = (cells[3] or '').strip()
        if ind.lower().startswith('header'):
            continue
        filled = sum(1 for c in cells if len((c or '').strip()) > 1)
        if filled < 3:
            continue
        out = (cells[0] or '').strip()
        if len(ind) >= 10 or len(out) >= 20 or filled >= 4:
            return r
    return 4


def _harvest_period_hint_excel(ws, data_start_row=4, max_pre_rows=40, max_cols=12):
    """Rows above the data block as text (for quarter/year), plus wide header scan."""
    parts = []
    last_pre = max(0, min((ws.max_row or 0), data_start_row - 1, max_pre_rows))
    for r in range(1, last_pre + 1):
        row_bits = []
        max_c = min(max_cols, ws.max_column or 0)
        for c in range(1, max_c + 1):
            row_bits.append(_cell_value_str(ws.cell(row=r, column=c).value))
        line = ' '.join(t for t in row_bits if t)
        if line:
            parts.append(line)
    return '\n'.join(parts)


def _ingest_excel_lpc_wide(ws, layout, office_name, quarter, year, extra_hint):
    """LPC / FYDP wide matrix: annual quantifiable target, accomplishment, MET/UNMET column."""
    hr = layout['header_row']
    cols = layout['cols']
    mc = layout['max_col']
    header_text = _harvest_period_hint_excel(ws, data_start_row=hr, max_pre_rows=55, max_cols=mc)
    if extra_hint:
        header_text = f'{extra_hint.strip()}\n{header_text}'
    auto_q, auto_y = _detect_quarter_year(header_text)
    # Annual rollups align with year-end reporting when quarter is not forced.
    current_q = quarter if quarter is not None else (auto_q or 4)
    current_y = year if year is not None else (auto_y or 2025)

    default_office, _ = Office.objects.get_or_create(name=office_name)
    dev_area = _harvest_lpc_development_area(ws, hr, mc)
    dev_area_current = canonicalize_opmm_dev_area_label(dev_area) if dev_area else ''

    oc, sc, pc, spc = (
        cols['outcome'],
        cols['strategy'],
        cols['program'],
        cols['sub_pap'],
    )
    ind_c = cols['indicator']
    ann_c = cols['annual_target']
    acmp_c = cols['accomplishment']
    status_c = cols['status']
    off_c = cols['office']

    last_out, last_strat, last_prog = '', '', ''
    last_ind_txt = ''
    seen_lpc_wide_ind_keys = set()
    data_start = hr + 1
    end_row = _excel_effective_max_row(ws, data_start)
    rows_written = 0
    empty_streak = 0

    for r in range(data_start, end_row + 1):
        # Some exports include multiple "Development Area:" sections in one sheet.
        # If we see a dev-area heading row, update the current dev area and reset
        # merged-header carry-overs so subsequent rows don't inherit the prior section.
        scan_cells = [
            _cell_value_str(ws.cell(row=r, column=c).value) for c in range(1, min(mc, 16) + 1)
        ]
        ind_txt = _ws_cell(ws, r, ind_c).strip()
        annual_txt = _ws_cell(ws, r, ann_c).strip()
        acmp_txt = _ws_cell(ws, r, acmp_c).strip()
        var_raw = _ws_cell(ws, r, status_c).strip() if status_c else ''
        out_txt_row = _ws_cell(ws, r, oc).strip() if oc else ''
        strat_txt_row = _ws_cell(ws, r, sc).strip() if sc else ''
        pap_txt_row = _ws_cell(ws, r, pc).strip() if pc else ''
        maybe_dev = _scan_row_for_development_area(scan_cells)
        if maybe_dev:
            dev_area_current = canonicalize_opmm_dev_area_label(maybe_dev)
            if _should_skip_row_as_dev_area_banner_only_lpc(
                scan_cells, ind_txt, annual_txt, acmp_txt, var_raw
            ) or _row_is_pure_opmm_section_banner_row(
                out_txt_row,
                strat_txt_row,
                pap_txt_row,
                ind_txt,
                acmp_txt,
                var_raw,
            ):
                last_out, last_strat, last_prog = '', '', ''
                last_ind_txt = ''
                continue

        if oc:
            t = _ws_cell(ws, r, oc).strip()
            if not t and r > data_start:
                t = _excel_vertical_merge_fill(ws, r - 1, oc, data_start)
            if t:
                if _matrix_label_changed(last_out, t):
                    last_strat, last_prog = '', ''
                last_out = t
        if sc:
            t = _ws_cell(ws, r, sc).strip()
            if not t and r > data_start:
                t = _excel_vertical_merge_fill(ws, r - 1, sc, data_start)
            if not t and r > data_start:
                t = _excel_cell_above(ws, r, sc)
            if t:
                if _matrix_label_changed(last_strat, t):
                    last_prog = ''
                last_strat = t
        if pc:
            raw_pc = _ws_cell(ws, r, pc).strip()
            t = raw_pc
            if not t and r > data_start:
                t = _excel_vertical_merge_fill(ws, r - 1, pc, data_start)
            if t:
                last_prog = t

        if ind_txt:
            last_ind_txt = ind_txt
        # Carry merged indicator cells only on continuation rows that still have KPI data.
        # Important: continuation rows must not reuse the *same* description as the anchor row,
        # or get_or_create(pap, description) + update_or_create(indicator, quarter, year) will
        # overwrite prior KPIs — the dashboard then shows far fewer indicators than the sheet.
        if ind_txt:
            use_ind = ind_txt
        elif last_ind_txt and (annual_txt or acmp_txt or var_raw):
            tail_a = (annual_txt or '').strip()
            tail_c = (acmp_txt or '').strip()
            tail_v = (var_raw or '').strip()
            if tail_a and tail_c:
                tail = f'{tail_a} · {tail_c[:120]}'
            elif tail_a:
                tail = tail_a
            elif tail_c:
                tail = tail_c[:200]
            elif tail_v:
                tail = tail_v[:120]
            else:
                tail = ''
            use_ind = f'{last_ind_txt} — {tail}' if tail else f'{last_ind_txt} — (LPC row {r})'
        else:
            use_ind = ''
        use_ind = use_ind.strip()

        if not use_ind and not annual_txt and not acmp_txt and not var_raw:
            empty_streak += 1
            if empty_streak >= 100:
                break
            continue
        empty_streak = 0

        if annual_txt and 'quantifiable' in _normalize_excel_header_text(annual_txt):
            continue
        if ind_txt and _normalize_excel_header_text(ind_txt) in (
            'performance indicators (based on the cascaded template)',
            'performance indicators (based on the submitted operational plan)',
        ):
            continue

        sub = _ws_cell(ws, r, spc).strip() if spc else ''
        papn = last_prog
        if sub:
            papn = f'{papn} — {sub}' if papn else sub

        raw_off = _ws_cell(ws, r, off_c).strip() if off_c else ''
        office_names = split_concerned_office_cell(raw_off)
        if office_names:
            row_offices = [
                Office.objects.get_or_create(name=n[:255])[0] for n in office_names
            ]
        else:
            row_offices = [default_office]

        out_db = last_out
        if dev_area_current:
            d_c = re.sub(r'[^a-z0-9]', '', dev_area_current.lower())
            t_c = re.sub(r'[^a-z0-9]', '', (last_out or '').lower())
            if last_out and d_c[:24] not in t_c:
                out_db = f'{dev_area_current} — {last_out}'
            elif not last_out:
                out_db = dev_area_current

        if (use_ind or '').strip():
            cur_ind = use_ind.strip()
            trial_key = (
                (papn or '').strip(),
                (out_db or '').strip(),
                (last_strat or '').strip(),
                cur_ind,
            )
            suf = 1
            while trial_key in seen_lpc_wide_ind_keys:
                suf += 1
                alt_ind = f'{cur_ind} · row {r}' if suf == 2 else f'{cur_ind} · row {r} ({suf})'
                trial_key = (
                    (papn or '').strip(),
                    (out_db or '').strip(),
                    (last_strat or '').strip(),
                    alt_ind,
                )
            if suf > 1:
                use_ind = trial_key[3]
            seen_lpc_wide_ind_keys.add(trial_key)

        exp = _parse_explicit_met_unmet_cell(var_raw) if var_raw else None
        ind_key = (
            use_ind
            if use_ind
            else (f'KPI ({annual_txt})' if annual_txt else f'LPC import row {r}')
        )

        for row_office in row_offices:
            _ingest_monitor_row_extended(
                row_office,
                out_db or 'Outcome',
                last_strat or 'Strategy',
                papn or 'Program / PAP',
                ind_key,
                annual_txt,
                acmp_txt,
                exp,
                current_q,
                current_y,
                variance_text=var_raw,
            )
            rows_written += 1
    return rows_written


def _ingest_excel_opmm_matrix(ws, layout, office_name, quarter, year, extra_hint):
    """OPMM monitoring matrix: sparse Outcome/Strategy/PAP columns, Actual + Variance (+ optional Remarks)."""
    hr = layout['header_row']
    cols = layout['cols']
    mc = layout['max_col']
    header_text = _harvest_period_hint_excel(ws, data_start_row=hr, max_pre_rows=55, max_cols=mc)
    if extra_hint:
        header_text = f'{extra_hint.strip()}\n{header_text}'
    auto_q, auto_y = _detect_quarter_year(header_text)
    current_q = quarter if quarter is not None else (auto_q or 4)
    current_y = year if year is not None else (auto_y or 2025)

    default_office, _ = Office.objects.get_or_create(name=office_name)
    dev_area = _harvest_lpc_development_area(ws, hr, mc)
    dev_area_current = canonicalize_opmm_dev_area_label(dev_area) if dev_area else ''

    oc, sc, pc, spc = (
        cols['outcome'],
        cols['strategy'],
        cols['program'],
        cols['sub_pap'],
    )
    ind_c = cols['indicator']
    acmp_c = cols['accomplishment']
    status_c = cols['status']
    remarks_c = cols['remarks']
    off_c = cols['office']

    last_out, last_strat, last_prog = '', '', ''
    last_ind_txt = ''
    seen_opmm_matrix_ind_keys = set()
    data_start = hr + 1
    end_row = _excel_effective_max_row(ws, data_start)
    rows_written = 0
    empty_streak = 0

    for r in range(data_start, end_row + 1):
        scan_cells = [
            _cell_value_str(ws.cell(row=r, column=c).value) for c in range(1, min(mc, 16) + 1)
        ]
        ind_txt = _ws_cell(ws, r, ind_c).strip()
        acmp_txt = _ws_cell(ws, r, acmp_c).strip()
        var_txt0 = _ws_cell(ws, r, status_c).strip() if status_c else ''
        rem_txt0 = _ws_cell(ws, r, remarks_c).strip() if remarks_c else ''
        out_txt_row = _ws_cell(ws, r, oc).strip() if oc else ''
        strat_txt_row = _ws_cell(ws, r, sc).strip() if sc else ''
        pap_txt_row = _ws_cell(ws, r, pc).strip() if pc else ''
        maybe_dev = _scan_row_for_development_area(scan_cells)
        if maybe_dev:
            dev_area_current = canonicalize_opmm_dev_area_label(maybe_dev)
            if _should_skip_row_as_dev_area_banner_only(
                scan_cells,
                ind_txt,
                acmp_txt,
                var_txt0,
                rem_txt0,
                outcome_text=out_txt_row,
                strategy_text=strat_txt_row,
                pap_text=pap_txt_row,
            ):
                last_out, last_strat, last_prog = '', '', ''
                last_ind_txt = ''
                continue

        if oc:
            t = _ws_cell(ws, r, oc).strip()
            if not t and r > data_start:
                t = _excel_vertical_merge_fill(ws, r - 1, oc, data_start)
            if t:
                if _matrix_label_changed(last_out, t):
                    last_strat, last_prog = '', ''
                last_out = t
        if sc:
            t = _ws_cell(ws, r, sc).strip()
            if not t and r > data_start:
                t = _excel_vertical_merge_fill(ws, r - 1, sc, data_start)
            if not t and r > data_start:
                t = _excel_cell_above(ws, r, sc)
            if t:
                if _matrix_label_changed(last_strat, t):
                    last_prog = ''
                last_strat = t
        if pc:
            raw_pc = _ws_cell(ws, r, pc).strip()
            t = raw_pc
            if not t and r > data_start:
                t = _excel_vertical_merge_fill(ws, r - 1, pc, data_start)
            if t:
                last_prog = t

        if ind_txt:
            last_ind_txt = ind_txt
        remarks_txt = _ws_cell(ws, r, remarks_c).strip() if remarks_c else ''

        if ind_txt:
            use_ind = ind_txt
        elif (acmp_txt or remarks_txt) and (last_out or last_strat or last_prog):
            # Same PAP block, new indicator row: do not merge into the previous indicator title.
            first_line = ''
            for block in (acmp_txt, remarks_txt):
                if not block:
                    continue
                line = block.strip().splitlines()[0].strip()
                if line:
                    first_line = line
                    break
            base = first_line if first_line else 'Performance indicator'
            use_ind = f'{base} (sheet row {r})'
        else:
            use_ind = ''
        use_ind = use_ind.strip()

        if not use_ind and not acmp_txt and not remarks_txt:
            empty_streak += 1
            if empty_streak >= 100:
                break
            continue
        empty_streak = 0

        if ind_txt and _normalize_excel_header_text(ind_txt) in (
            'performance indicators (based on the cascaded template)',
            'performance indicators (based on the submitted operational plan)',
        ):
            continue

        sub = _ws_cell(ws, r, spc).strip() if spc else ''
        papn = last_prog
        if sub:
            papn = f'{papn} — {sub}' if papn else sub

        raw_off = _ws_cell(ws, r, off_c).strip() if off_c else ''
        office_names = split_concerned_office_cell(raw_off)
        if office_names:
            row_offices = [
                Office.objects.get_or_create(name=n[:255])[0] for n in office_names
            ]
        else:
            row_offices = [default_office]

        out_db = last_out
        if dev_area_current:
            d_c = re.sub(r'[^a-z0-9]', '', dev_area_current.lower())
            t_c = re.sub(r'[^a-z0-9]', '', (last_out or '').lower())
            if last_out and d_c[:24] not in t_c:
                out_db = f'{dev_area_current} — {last_out}'
            elif not last_out:
                out_db = dev_area_current

        cur_ind = (use_ind or '').strip()
        trial_key = (
            (papn or '').strip(),
            (out_db or '').strip(),
            (last_strat or '').strip(),
            cur_ind,
        )
        suf = 1
        while trial_key in seen_opmm_matrix_ind_keys:
            suf += 1
            alt_ind = f'{cur_ind} · row {r}' if suf == 2 else f'{cur_ind} · row {r} ({suf})'
            trial_key = (
                (papn or '').strip(),
                (out_db or '').strip(),
                (last_strat or '').strip(),
                alt_ind,
            )
        if suf > 1:
            use_ind = trial_key[3]
        seen_opmm_matrix_ind_keys.add(trial_key)

        var_raw = _ws_cell(ws, r, status_c).strip() if status_c else ''
        exp = _parse_explicit_met_unmet_cell(var_raw) if var_raw else None
        ind_key = use_ind if use_ind else (f'KPI ({acmp_txt})' if acmp_txt else f'OPMM matrix row {r}')

        full_acmp = acmp_txt
        if remarks_txt:
            if full_acmp:
                full_acmp = f'{full_acmp}\n\n— Remarks —\n{remarks_txt}'
            else:
                full_acmp = remarks_txt

        for row_office in row_offices:
            _ingest_monitor_row_extended(
                row_office,
                out_db or 'Outcome',
                last_strat or 'Strategy',
                papn or 'Program / PAP',
                ind_key,
                ind_key,
                full_acmp,
                exp,
                current_q,
                current_y,
                variance_text=var_raw,
            )
            rows_written += 1
    return rows_written


def ingest_excel_monitor(file_path, office_name="OVCRDES", quarter=None, year=None, extra_hint=''):
    """Parse an Excel monitor workbook (.xlsx).

    Supports:

    * **Narrow** A–E (active sheet): outcome, strategy, PAP, indicator (with optional target text),
      accomplishment; optional **F** = office name.
    * **LPC / FYDP wide** on **each worksheet** that matches the layout (merged outcome/strategy/program,
      annual quantifiable target, **Actual Accomplishments** or quarterly accomplishment columns,
      and **Status of Accomplishment** or **Variance** for MET/UNMET / +N variance).
    * **OPMM monitoring matrix** (e.g. BatStateU OVCRDES): numbered columns, Actual Accomplishments + Variance,
      optional Remarks; **no** annual quantifiable target column — sparse parent columns carried down.

    Returns the number of data rows written (performance rows upserted).
    """
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=False, data_only=True)
    rows_written = 0
    try:
        # Many institutional OPMM workbooks use one LPC-wide block per worksheet
        # (e.g. each development area). Ingest every sheet that matches the layout.
        total_from_sheets = 0
        for ws in wb.worksheets:
            layout = _detect_lpc_wide_layout(ws)
            if layout:
                total_from_sheets += _ingest_excel_lpc_wide(
                    ws, layout, office_name, quarter, year, extra_hint
                )
            else:
                mlayout = _detect_opmm_matrix_layout(ws)
                if mlayout:
                    total_from_sheets += _ingest_excel_opmm_matrix(
                        ws, mlayout, office_name, quarter, year, extra_hint
                    )
        if total_from_sheets:
            return total_from_sheets

        ws = wb.active
        data_start = _excel_data_start_row(ws)
        header_text = _harvest_period_hint_excel(ws, data_start_row=data_start)
        if extra_hint:
            header_text = f'{extra_hint.strip()}\n{header_text}'
        auto_q, auto_y = _detect_quarter_year(header_text)
        current_q = quarter if quarter is not None else (auto_q or 1)
        current_y = year if year is not None else (auto_y or 2025)

        default_office, _ = Office.objects.get_or_create(name=office_name)

        max_col = ws.max_column or 0
        end_row = _excel_effective_max_row(ws, data_start)
        empty_streak = 0
        dev_area_current = None
        for r in range(data_start, end_row + 1):
            cells = [
                _cell_value_str(ws.cell(row=r, column=c).value) for c in range(1, 6)
            ]
            if len(cells) < 5:
                continue
            if not any(cells):
                empty_streak += 1
                if empty_streak >= 100:
                    break
                continue
            empty_streak = 0
            if _excel_row_looks_like_header(cells):
                continue
            if (cells[3] or '').strip().lower().startswith('header'):
                continue
            maybe_dev = _scan_row_for_development_area(cells)
            if maybe_dev:
                dev_area_current = canonicalize_opmm_dev_area_label(maybe_dev)
                if _should_skip_row_as_dev_area_banner_only(
                    cells,
                    cells[3] if len(cells) > 3 else '',
                    cells[4] if len(cells) > 4 else '',
                    '',
                    '',
                    outcome_text=cells[0] if cells else '',
                    strategy_text=cells[1] if len(cells) > 1 else '',
                    pap_text=cells[2] if len(cells) > 2 else '',
                ):
                    continue
            if not any(
                [
                    (cells[0] or '').strip(),
                    (cells[1] or '').strip(),
                    (cells[2] or '').strip(),
                    (cells[3] or '').strip(),
                ]
            ):
                continue
            raw_f = _cell_value_str(ws.cell(row=r, column=6).value) if max_col >= 6 else ''
            office_names = split_concerned_office_cell(raw_f)
            if office_names:
                row_offices = [
                    Office.objects.get_or_create(name=n[:255])[0] for n in office_names
                ]
            else:
                row_offices = [default_office]
            for row_office in row_offices:
                # Prefix the outcome with the most recent development-area heading so
                # the performance viewer/admin can consistently bucket records.
                if dev_area_current:
                    out_txt = (cells[0] or '').strip()
                    d_c = re.sub(r'[^a-z0-9]', '', dev_area_current.lower())
                    o_c = re.sub(r'[^a-z0-9]', '', out_txt.lower())
                    if out_txt and d_c[:24] in o_c:
                        pass
                    else:
                        cells = list(cells)
                        cells[0] = f'{dev_area_current} — {out_txt or "Outcome"}'
                _ingest_monitor_row(row_office, cells, current_q, current_y)
                rows_written += 1
    finally:
        wb.close()
    return rows_written

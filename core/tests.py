import os
import re
import sqlite3
import tempfile
from pathlib import Path
from unittest.mock import MagicMock

from django.conf import settings
from django.core.management import call_command
from django.core.management.base import CommandError
from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client, TestCase, TransactionTestCase, override_settings
from django.urls import reverse
from docx import Document

from .models import ActivityLog, Announcement, AnnouncementRead, Indicator, Office, PerformanceRecord, StrategicLevel
from .services import (
    _detect_quarter_year,
    _extract_opmm_section_banner,
    _explicit_met_unmet_from_accomplishment_tail,
    _parse_explicit_met_unmet_cell,
    _row_is_pure_opmm_section_banner_row,
    extract_actual_for_monitor,
    extract_number,
    extract_target_for_monitor,
    ingest_excel_monitor,
    record_is_pillar_banner_only,
    ingest_word_monitor,
)
from .views import (
    _analytics_row_tooltip_label_for_pillar,
    _build_office_chart_payload,
    _search_suggestion_focus,
    _validate_sqlite_backup_path,
)


def _build_min_monitor_docx(path):
    """Minimal .docx matching ingest_word_monitor table layout (data from row index 3)."""
    doc = Document()
    doc.add_paragraph('Status Report 2nd Quarter 2026')
    table = doc.add_table(rows=4, cols=5)
    for ri in range(3):
        for ci in range(5):
            table.rows[ri].cells[ci].text = f'header_{ri}_{ci}'
    row = table.rows[3].cells
    row[0].text = 'Academic Leadership and excellence'
    row[1].text = 'Strategy Alpha'
    row[2].text = 'PAP Sample'
    row[3].text = 'Indicator target (100)'
    row[4].text = 'Accomplishment (95)'
    doc.save(path)


def _build_word_opmm_hash_header_docx(path):
    """OPMM-style table with leading # column so variance is not at index 5."""
    doc = Document()
    doc.add_paragraph('Status Report 2nd Quarter 2026')
    table = doc.add_table(rows=5, cols=8)
    for ri in range(2):
        for ci in range(8):
            table.rows[ri].cells[ci].text = f'pre_{ri}_{ci}'
    hdr = [
        '#',
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    r = table.rows[4].cells
    r[0].text = '1'
    r[1].text = 'Outcome One'
    r[2].text = 'Strategy Two'
    r[3].text = 'PAP Three'
    r[4].text = 'Graduation target (100)'
    r[5].text = 'Narrative says eighty-eight (88)'
    r[6].text = 'MET'
    r[7].text = 'OK'
    doc.save(path)


def _build_word_merged_continuation_row_docx(path):
    """Second body row has empty merged Outcome/Strategy/PAP cells (common in Word OPMM tables)."""
    doc = Document()
    doc.add_paragraph('Operational Plan Monitoring Matrix 1st Quarter 2024')
    table = doc.add_table(rows=5, cols=7)
    for ri in range(2):
        for ci in range(7):
            table.rows[ri].cells[ci].text = f'pre_{ri}_{ci}'
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    r3 = table.rows[3].cells
    r3[0].text = 'Outcome Shared Meta'
    r3[1].text = 'Strategy First Stream'
    r3[2].text = 'PAP Young Program'
    r3[3].text = 'Indicator students count (10)'
    r3[4].text = 'Narrative about students for row one.'
    r3[5].text = 'MET'
    r3[6].text = ''
    r4 = table.rows[4].cells
    r4[0].text = ''
    r4[1].text = ''
    r4[2].text = ''
    r4[3].text = 'Indicator sessions count (5)'
    r4[4].text = 'Narrative about sessions for row two.'
    r4[5].text = 'N/A'
    r4[6].text = ''
    doc.save(path)


def _build_word_matrix_header_row0_four_kpis_docx(path):
    """OPMM header on row 0 with KPIs on rows 1–4 (regression: data_start must not skip rows 1–2)."""
    doc = Document()
    doc.add_paragraph('OPMM Research and Innovation FY 2024')
    table = doc.add_table(rows=5, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[0].cells[ci].text = h
    kpi_rows = [
        (
            'RESEARCH AND INNOVATION: Engineering solutions',
            'Strategy R1',
            'PAP R1',
            'Publications in high-impact journals target (10)',
            'Narrative accomplishment text for KPI one exceeds forty characters total.',
            'MET',
            '',
        ),
        (
            'RESEARCH AND INNOVATION: Engineering solutions',
            'Strategy R1',
            'PAP R1',
            'External research funding target (50)',
            'Second KPI narrative accomplishment text is also longer than forty chars.',
            'UNMET',
            '',
        ),
        (
            'RESEARCH AND INNOVATION: Engineering solutions',
            'Strategy R1',
            'PAP R1',
            'Patents filed target (4)',
            'Third KPI narrative accomplishment text is longer than forty characters here.',
            'MET',
            '',
        ),
        (
            'RESEARCH AND INNOVATION: Engineering solutions',
            'Strategy R1',
            'PAP R1',
            'Innovation workshops target (12)',
            'Fourth KPI narrative accomplishment text exceeds forty characters as required.',
            'MET',
            '',
        ),
    ]
    for ri, cells in enumerate(kpi_rows, start=1):
        for ci, val in enumerate(cells):
            table.rows[ri].cells[ci].text = val
    doc.save(path)


def _build_word_social_banner_row_plus_two_kpis_docx(path):
    """Pillar title repeated across columns + MET (must not ingest); then two real KPI rows."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2026')
    table = doc.add_table(rows=6, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    banner = 'SOCIAL RESPONSIBILITY: Engineering Pathways for Families and Communities'
    rb = table.rows[3].cells
    for j in range(4):
        rb[j].text = banner
    rb[4].text = '—'
    rb[5].text = 'MET'
    rb[6].text = ''
    sr_out = banner
    r4 = table.rows[4].cells
    r4[0].text = sr_out
    r4[1].text = 'Strategy for community engagement'
    r4[2].text = 'Engineering Pathways PAP'
    r4[3].text = 'Number of community outreach programs conducted (12)'
    r4[4].text = 'Narrative for outreach programs exceeds forty characters minimum length.'
    r4[5].text = 'MET'
    r5 = table.rows[5].cells
    r5[0].text = sr_out
    r5[1].text = 'Strategy for community engagement'
    r5[2].text = 'Engineering Pathways PAP'
    r5[3].text = 'Number of families served through engineering pathways (50)'
    r5[4].text = 'Narrative for families served exceeds forty characters minimum length.'
    r5[5].text = 'UNMET'
    doc.save(path)


def _build_word_ri_banner_young_lifters_writeshop_docx(path):
    """R&I pillar banner row + two KPIs (Young LIFTERS then Writeshop), matching OVCRDES Word layout."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
    table = doc.add_table(rows=6, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    pillar = 'RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development'
    rb = table.rows[3].cells
    for j in range(7):
        rb[j].text = pillar if j < 4 else ''
    rb[5].text = ''
    rb[6].text = ''
    r4 = table.rows[4].cells
    r4[0].text = 'University Research and Innovation Ecosystem Advanced'
    r4[1].text = 'Promote cross-disciplinary and interdisciplinary research collaborations'
    r4[2].text = 'Young Leadership and Innovation for Future Thinkers (Young LIFTERS) Program'
    r4[3].text = 'Number of students participated in the program'
    r4[4].text = (
        'One (1) implemented monitoring of research projects of YLP mentees '
        'with sufficient narrative length for ingest.'
    )
    r4[5].text = 'MET'
    r4[6].text = 'RMS conducted monitoring and assessed the progress of YLP research project.'
    r5 = table.rows[5].cells
    r5[0].text = ''
    r5[1].text = 'Strengthen capability building of researchers in different fields'
    r5[2].text = 'Regular Writeshops Kapehan and HUNTahan Session'
    r5[3].text = 'Number of conducted writeshop/HUNTahan Session'
    r5[4].text = 'Zero (0) regular Writeshops, Kapehan and HUNTahan Session'
    r5[5].text = 'N/A'
    r5[6].text = (
        'The target was initially set for the first quarter; funds may move to Q2. '
        'Best Practices: offer relevant training and support for faculty research.'
    )
    doc.save(path)


def _build_word_ri_first_kpi_pillar_cloned_in_hierarchy_docx(path):
    """Word merge: pillar title cloned into Outcome–Indicator cells; accomplishment is still KPI data."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
    table = doc.add_table(rows=5, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    pillar = 'RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development'
    r4 = table.rows[3].cells
    for j in range(4):
        r4[j].text = pillar
    r4[4].text = (
        'One (1) implemented monitoring of research projects of YLP mentees '
        'with sufficient narrative length for ingest.'
    )
    r4[5].text = 'MET'
    r5 = table.rows[4].cells
    r5[0].text = ''
    r5[1].text = 'Strengthen capability building of researchers in different fields'
    r5[2].text = 'Regular Writeshops Kapehan and HUNTahan Session'
    r5[3].text = 'Number of conducted writeshop/HUNTahan Session'
    r5[4].text = 'Zero (0) regular Writeshops, Kapehan and HUNTahan Session'
    r5[5].text = 'N/A'
    doc.save(path)


def _build_word_ri_first_kpi_duplicate_outcome_strategy_docx(path):
    """Word merge quirk: Outcome text duplicated in Strategy on the first KPI row (must still ingest)."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
    table = doc.add_table(rows=5, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    uni = 'University Research and Innovation Ecosystem Advanced'
    r4 = table.rows[3].cells
    r4[0].text = uni
    r4[1].text = uni
    r4[2].text = 'Young Leadership and Innovation for Future Thinkers (Young LIFTERS) Program'
    r4[3].text = 'Number of students participated in the program'
    r4[4].text = (
        'One (1) implemented monitoring of research projects of YLP mentees '
        'with sufficient narrative length for ingest.'
    )
    r4[5].text = 'MET'
    r5 = table.rows[4].cells
    r5[0].text = ''
    r5[1].text = 'Strengthen capability building of researchers in different fields'
    r5[2].text = 'Regular Writeshops Kapehan and HUNTahan Session'
    r5[3].text = 'Number of conducted writeshop/HUNTahan Session'
    r5[4].text = 'Zero (0) regular Writeshops, Kapehan and HUNTahan Session'
    r5[5].text = 'N/A'
    doc.save(path)


def _build_word_ovcrdes_ri_four_kpis_merged_blanks_docx(path):
    """OVCRDES-style R&I block: two PAPs × two KPI rows; rows 4 and 6 have blank merged hierarchy."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
    table = doc.add_table(rows=7, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    ri_out = (
        'RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development'
    )
    st = 'Strengthen campus research niching'
    pap_y = 'Young LIFTers Program for targeted students'
    pap_f = 'Fintech Applications Development Program for BatStateU Lipa'
    r3 = table.rows[3].cells
    r3[0].text, r3[1].text, r3[2].text = ri_out, st, pap_y
    r3[3].text = 'Number of students participated in the program (100)'
    r3[4].text = 'Narrative for students KPI exceeds forty characters minimum length.'
    r3[5].text = 'MET'
    r4 = table.rows[4].cells
    for j in range(3):
        r4[j].text = ''
    r4[3].text = 'Number of conducted writeshop HUNTahan Session (10)'
    r4[4].text = 'Narrative for writeshop KPI exceeds forty characters minimum length.'
    r4[5].text = 'MET'
    r5 = table.rows[5].cells
    r5[0].text, r5[1].text, r5[2].text = ri_out, st, pap_f
    r5[3].text = 'Number of capacity-building programs conducted focused on research publication (8)'
    r5[4].text = 'Narrative for capacity building KPI exceeds forty characters minimum length.'
    r5[5].text = 'MET'
    r6 = table.rows[6].cells
    for j in range(3):
        r6[j].text = ''
    r6[3].text = (
        'Number of Concept Proposal Development and Ideation for the Fintech Research Center conducted (6)'
    )
    r6[4].text = 'Narrative for concept proposal KPI exceeds forty characters minimum length.'
    r6[5].text = 'MET'
    doc.save(path)


def _build_word_ovcrdes_ri_four_kpis_empty_indicator_continuation_docx(path):
    """Like merged hierarchy blanks, but indicator column is empty on continuation rows (Word merge)."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
    table = doc.add_table(rows=7, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    ri_out = (
        'RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development'
    )
    st = 'Strengthen campus research niching'
    pap_y = 'Young LIFTers Program for targeted students'
    pap_f = 'Fintech Applications Development Program for BatStateU Lipa'
    r3 = table.rows[3].cells
    r3[0].text, r3[1].text, r3[2].text = ri_out, st, pap_y
    r3[3].text = 'Number of students participated in the program (100)'
    r3[4].text = 'Narrative for students KPI exceeds forty characters minimum length.'
    r3[5].text = 'MET'
    r4 = table.rows[4].cells
    for j in range(3):
        r4[j].text = ''
    r4[3].text = ''
    r4[4].text = 'Narrative for writeshop KPI exceeds forty characters minimum length.'
    r4[5].text = 'MET'
    r5 = table.rows[5].cells
    r5[0].text, r5[1].text, r5[2].text = ri_out, st, pap_f
    r5[3].text = 'Number of capacity-building programs conducted focused on research publication (8)'
    r5[4].text = 'Narrative for capacity building KPI exceeds forty characters minimum length.'
    r5[5].text = 'MET'
    r6 = table.rows[6].cells
    for j in range(3):
        r6[j].text = ''
    r6[3].text = ''
    r6[4].text = 'Narrative for concept proposal KPI exceeds forty characters minimum length.'
    r6[5].text = 'MET'
    doc.save(path)


def _build_word_ovcrdes_ri_four_kpis_empty_indicator_short_narrative_docx(path):
    """Empty indicator on continuation rows with accomplishment text under 18 chars (strict-KPI false)."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
    table = doc.add_table(rows=7, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    ri_out = (
        'RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development'
    )
    st = 'Strengthen campus research niching'
    pap_y = 'Young LIFTers Program for targeted students'
    pap_f = 'Fintech Applications Development Program for BatStateU Lipa'
    r3 = table.rows[3].cells
    r3[0].text, r3[1].text, r3[2].text = ri_out, st, pap_y
    r3[3].text = 'Number of students participated in the program (100)'
    r3[4].text = 'Narrative for students KPI exceeds forty characters minimum length.'
    r3[5].text = 'MET'
    r4 = table.rows[4].cells
    for j in range(3):
        r4[j].text = ''
    r4[3].text = ''
    r4[4].text = 'Short ac'
    r4[5].text = 'MET'
    r5 = table.rows[5].cells
    r5[0].text, r5[1].text, r5[2].text = ri_out, st, pap_f
    r5[3].text = 'Number of capacity-building programs conducted focused on research publication (8)'
    r5[4].text = 'Narrative for capacity building KPI exceeds forty characters minimum length.'
    r5[5].text = 'MET'
    r6 = table.rows[6].cells
    for j in range(3):
        r6[j].text = ''
    r6[3].text = ''
    r6[4].text = 'Tiny'
    r6[5].text = 'MET'
    doc.save(path)


def _build_word_opmm_duplicate_indicator_column_two_rows_docx(path):
    """Two body rows with identical indicator cell text (Word merge repeat); accomplishments differ."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 2nd Quarter 2026')
    table = doc.add_table(rows=5, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    row_a = (
        'OUTCOME ONE',
        'Strategy one',
        'PAP one',
        'Shared KPI title with target (99)',
        'First narrative accomplishment is longer than forty characters as required here.',
        'MET',
        '',
    )
    row_b = (
        'OUTCOME ONE',
        'Strategy one',
        'PAP one',
        'Shared KPI title with target (99)',
        'Second narrative accomplishment is also longer than forty characters for uniqueness.',
        'UNMET',
        '',
    )
    for ri, cells in enumerate([row_a, row_b], start=3):
        for ci, val in enumerate(cells):
            table.rows[ri].cells[ci].text = val
    doc.save(path)


def _build_word_opmm_duplicate_indicator_cell_four_rows_docx(path):
    """Four body rows share one merged indicator label; accomplishments differ (4 → 3 regression)."""
    doc = Document()
    doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 2nd Quarter 2026')
    table = doc.add_table(rows=7, cols=7)
    hdr = [
        'Outcome (2)',
        'Strategy/ies (3)',
        'PAP (4)',
        'Performance Indicator/s (5)',
        'Actual Accomplishments (6)',
        'Variance (7)',
        'Remarks (8)',
    ]
    for ci, h in enumerate(hdr):
        table.rows[2].cells[ci].text = h
    shared = 'Shared KPI title with target (99)'
    narratives = [
        'First narrative accomplishment is longer than forty characters as required here.',
        'Second narrative accomplishment is also longer than forty characters for uniqueness.',
        'Third narrative accomplishment is also longer than forty characters for uniqueness.',
        'Fourth narrative accomplishment is also longer than forty characters for uniqueness.',
    ]
    stats = ['MET', 'UNMET', 'MET', 'MET']
    for k in range(4):
        row = (
            'OUTCOME ONE',
            'Strategy one',
            'PAP one',
            shared,
            narratives[k],
            stats[k],
            '',
        )
        for ci, val in enumerate(row):
            table.rows[3 + k].cells[ci].text = val
    doc.save(path)


def _build_two_table_monitor_docx(path):
    """Two monitor tables in one file; ingest should read both."""
    doc = Document()
    doc.add_paragraph('Status Report 2nd Quarter 2026')

    t1 = doc.add_table(rows=5, cols=5)
    for ri in range(2):
        for ci in range(5):
            t1.rows[ri].cells[ci].text = f't1_header_{ri}_{ci}'
    t1.rows[2].cells[0].text = 'Outcome'
    t1.rows[2].cells[1].text = 'Strategy'
    t1.rows[2].cells[2].text = 'PAP'
    t1.rows[2].cells[3].text = 'Indicator/s'
    t1.rows[2].cells[4].text = 'Actual Accomplishments'
    row1 = t1.rows[4].cells
    row1[0].text = 'Outcome One'
    row1[1].text = 'Strategy A'
    row1[2].text = 'PAP A'
    row1[3].text = 'Indicator A (10)'
    row1[4].text = 'Accomplishment A (8)'

    t2 = doc.add_table(rows=5, cols=5)
    for ri in range(2):
        for ci in range(5):
            t2.rows[ri].cells[ci].text = f't2_header_{ri}_{ci}'
    t2.rows[2].cells[0].text = 'Outcome'
    t2.rows[2].cells[1].text = 'Strategy'
    t2.rows[2].cells[2].text = 'PAP'
    t2.rows[2].cells[3].text = 'Indicator/s'
    t2.rows[2].cells[4].text = 'Actual Accomplishments'
    row2 = t2.rows[4].cells
    row2[0].text = 'Outcome Two'
    row2[1].text = 'Strategy B'
    row2[2].text = 'PAP B'
    row2[3].text = 'Indicator B (5)'
    row2[4].text = 'Accomplishment B (5)'

    doc.save(path)


def _build_min_monitor_xlsx(path):
    """Minimal .xlsx matching ingest_excel_monitor layout (rows 1–3 header, data from row 4)."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    for ri in range(3):
        for ci in range(5):
            v = f'header_{ri}_{ci}'
            if ri == 0 and ci == 0:
                v = 'Status Report 4th Quarter 2025'
            ws.cell(row=ri + 1, column=ci + 1, value=v)
    ws.cell(row=4, column=1, value='Academic Leadership and excellence')
    ws.cell(row=4, column=2, value='Strategy Alpha')
    ws.cell(row=4, column=3, value='PAP Sample')
    ws.cell(row=4, column=4, value='Indicator target (100)')
    ws.cell(row=4, column=5, value='Accomplishment (95)')
    wb.save(path)


class ExtractNumberTests(TestCase):
    def test_parentheses_preferred(self):
        self.assertEqual(extract_number('Target (88)'), 88.0)

    def test_first_digit_group(self):
        self.assertEqual(extract_number('Score 42 items'), 42.0)

    def test_empty(self):
        self.assertEqual(extract_number(''), 0.0)
        self.assertEqual(extract_number(None), 0.0)


class MonitorExtractTests(TestCase):
    def test_target_ignores_fy_year_without_real_kpi(self):
        self.assertIsNone(
            extract_target_for_monitor(
                'Local APRC for FY 2025 accomplishments facilitated'
            )
        )

    def test_target_from_parentheses(self):
        self.assertEqual(extract_target_for_monitor('Graduation rate target (85)'), 85.0)

    def test_target_skips_only_year_paren_then_finds_goal(self):
        self.assertEqual(
            extract_target_for_monitor('Note (2025) service level (12) units'),
            12.0,
        )

    def test_actual_accepts_small_integer(self):
        self.assertEqual(extract_actual_for_monitor('1'), 1.0)
        self.assertEqual(extract_actual_for_monitor('88% achieved'), 88.0)

    def test_opmm_section_banner_detects_pillar_without_development_area_prefix(self):
        s = _extract_opmm_section_banner(
            'SOCIAL RESPONSIBILITY: Engineering Pathways for Families and Communities'
        )
        self.assertIsNotNone(s)
        self.assertRegex(s, r'(?i)social\s+responsibility')
        self.assertRegex(s, r'(?i)engineering\s+pathways')
        t = _extract_opmm_section_banner(
            'Table 4. RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development'
        )
        self.assertIsNotNone(t)
        self.assertRegex(t, r'(?i)research\s+and\s+innovation')

    def test_pure_section_banner_row_detects_repeated_pillar_columns(self):
        banner = 'SOCIAL RESPONSIBILITY: Engineering Pathways for Families and Communities'
        self.assertTrue(
            _row_is_pure_opmm_section_banner_row(banner, banner, banner, banner, '—', 'MET', '')
        )
        self.assertFalse(
            _row_is_pure_opmm_section_banner_row(
                banner,
                'Strategy for community engagement',
                'Engineering Pathways PAP',
                'Number of community outreach programs conducted (12)',
                'Narrative for outreach programs exceeds forty characters minimum length.',
                'MET',
                '',
            )
        )


class VarianceMetUnmetRuleTests(TestCase):
    """Numeric variance: ``>= 0`` is MET, anything below 0 is UNMET (plain ``0`` is MET)."""

    def test_zero_variance_is_met(self):
        self.assertEqual(_parse_explicit_met_unmet_cell('0'), 'MET')
        self.assertEqual(_parse_explicit_met_unmet_cell(' 0 '), 'MET')
        self.assertEqual(_parse_explicit_met_unmet_cell('+0'), 'MET')
        self.assertEqual(_parse_explicit_met_unmet_cell('-0'), 'MET')
        self.assertEqual(_parse_explicit_met_unmet_cell('0%'), 'MET')

    def test_positive_variance_is_met(self):
        self.assertEqual(_parse_explicit_met_unmet_cell('5'), 'MET')
        self.assertEqual(_parse_explicit_met_unmet_cell('+5'), 'MET')
        self.assertEqual(_parse_explicit_met_unmet_cell('+12.5'), 'MET')
        self.assertEqual(_parse_explicit_met_unmet_cell('+5%'), 'MET')

    def test_negative_variance_is_unmet(self):
        self.assertEqual(_parse_explicit_met_unmet_cell('-1'), 'UNMET')
        self.assertEqual(_parse_explicit_met_unmet_cell('-12'), 'UNMET')
        self.assertEqual(_parse_explicit_met_unmet_cell('-3.5'), 'UNMET')
        self.assertEqual(_parse_explicit_met_unmet_cell('-3%'), 'UNMET')

    def test_record_status_uses_variance_over_numeric_target(self):
        office = Office.objects.create(name='VarianceRuleOffice')
        out = StrategicLevel.objects.create(name='Outcome', level_type='OUTCOME')
        st = StrategicLevel.objects.create(name='Strategy', level_type='STRATEGY', parent=out)
        pap = StrategicLevel.objects.create(
            name='PAP', level_type='PAP', parent=st, office=office
        )
        ind = Indicator.objects.create(pap=pap, description='Variance rule indicator')
        rec_zero = PerformanceRecord.objects.create(
            indicator=ind, quarter=1, year=2026, target_value=10, actual_value=10,
            variance_text='0',
        )
        self.assertEqual(rec_zero.status, 'MET')
        rec_zero.delete()
        rec_neg = PerformanceRecord.objects.create(
            indicator=ind, quarter=1, year=2026, target_value=10, actual_value=8,
            variance_text='-2',
        )
        self.assertEqual(rec_neg.status, 'UNMET')


class ViewerDevAreaClassifyTests(TestCase):
    """Regression: pillar matching order must not hide R&I rows that mention ``extension`` in text."""

    def test_research_and_innovation_wins_when_blob_contains_extension_substring(self):
        from core.views import _viewer_dev_area_key_from_compact_text

        blob = 'researchandinnovationengineeringprogramdeadlineextensiongranted'
        self.assertEqual(
            _viewer_dev_area_key_from_compact_text(blob),
            'research_and_innovation',
        )


class ViewerOutcomeFirstBucketTests(TestCase):
    def test_ri_outcome_beats_extension_word_in_indicator_text(self):
        from core.views import _viewer_record_dev_area_key

        office = Office.objects.create(name='OfficeViewerRiExtTest')
        out = StrategicLevel.objects.create(
            name='RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development',
            level_type='OUTCOME',
        )
        st = StrategicLevel.objects.create(name='Strategy line', level_type='STRATEGY', parent=out)
        pap = StrategicLevel.objects.create(name='PAP line', level_type='PAP', parent=st, office=office)
        ind = Indicator.objects.create(
            pap=pap,
            description='Please approve a one-week extension for reporting (deadline note).',
        )
        rec = PerformanceRecord.objects.create(indicator=ind, quarter=2, year=2026)
        self.assertEqual(_viewer_record_dev_area_key(rec), 'research_and_innovation')


class DetectQuarterYearTests(TestCase):
    def test_second_quarter_and_year(self):
        q, y = _detect_quarter_year('Report Second Quarter 2026')
        self.assertEqual(q, 2)
        self.assertEqual(y, 2026)

    def test_q_notation(self):
        q, y = _detect_quarter_year('Overview Q3 2025')
        self.assertEqual(q, 3)
        self.assertEqual(y, 2025)

    def test_lpc_style_filename(self):
        q, y = _detect_quarter_year('Copy of LPC-OPMM-Q4-2025.xlsx')
        self.assertEqual(q, 4)
        self.assertEqual(y, 2025)

    def test_year_prefers_trailing_calendar_year(self):
        q, y = _detect_quarter_year('Baseline 2024 notes Overview Q3 2025 wrap-up')
        self.assertEqual(q, 3)
        self.assertEqual(y, 2025)


class IngestWordMonitorTests(TestCase):
    def test_ingest_creates_performance_row(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_min_monitor_docx(path)
            ingest_word_monitor(path, office_name='UnitTestOffice', quarter=2, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertTrue(Office.objects.filter(name='UnitTestOffice').exists())
        rec = PerformanceRecord.objects.filter(quarter=2, year=2026).first()
        self.assertIsNotNone(rec)
        self.assertEqual(rec.target_value, 100.0)
        self.assertEqual(rec.actual_value, 95.0)
        self.assertEqual(rec.status, 'UNMET')

    def test_word_hash_column_maps_variance_met(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_opmm_hash_header_docx(path)
            ingest_word_monitor(path, office_name='HashColOffice', quarter=2, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        rec = PerformanceRecord.objects.filter(
            quarter=2, year=2026, indicator__pap__office__name='HashColOffice'
        ).first()
        self.assertIsNotNone(rec)
        self.assertEqual((rec.variance_text or '').strip().upper(), 'MET')
        self.assertEqual(rec.explicit_status, 'MET')
        self.assertEqual(rec.status, 'MET')

    def test_explicit_status_from_accomplishment_tail(self):
        self.assertEqual(_explicit_met_unmet_from_accomplishment_tail('done\nMET'), 'MET')
        self.assertEqual(_explicit_met_unmet_from_accomplishment_tail('x\n\nUNMET'), 'UNMET')
        self.assertIsNone(_explicit_met_unmet_from_accomplishment_tail('partially met target'))

    def test_ingest_reads_multiple_word_tables(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_two_table_monitor_docx(path)
            written = ingest_word_monitor(path, office_name='MultiTableOffice', quarter=2, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(written, 2)
        recs = PerformanceRecord.objects.filter(
            indicator__pap__office__name='MultiTableOffice',
            quarter=2,
            year=2026,
        )
        self.assertEqual(recs.count(), 2)

    def test_word_second_row_blank_merge_cells_inherit_hierarchy(self):
        """Merged blank hierarchy on row 2 must not drop the second KPI onto a stray PAP block."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_merged_continuation_row_docx(path)
            ingest_word_monitor(path, office_name='MergeWordOffice', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        recs = list(
            PerformanceRecord.objects.filter(
                quarter=1, year=2024, indicator__pap__office__name='MergeWordOffice'
            ).select_related('indicator', 'indicator__pap')
        )
        self.assertEqual(len(recs), 2)
        pap_names = {r.indicator.pap.name for r in recs}
        self.assertEqual(pap_names, {'PAP Young Program'})
        descs = {r.indicator.description for r in recs}
        self.assertTrue(any('students' in d.lower() for d in descs))
        self.assertTrue(any('sessions' in d.lower() for d in descs))

    def test_word_header_row_zero_ingests_all_kpis(self):
        """When the matrix header is row 0, KPIs on rows 1–4 must all ingest (not only rows 3–4)."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_matrix_header_row0_four_kpis_docx(path)
            n = ingest_word_monitor(path, office_name='EarlyHeaderOffice', quarter=4, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 4)
        recs = PerformanceRecord.objects.filter(
            quarter=4, year=2024, indicator__pap__office__name='EarlyHeaderOffice'
        ).select_related('indicator')
        self.assertEqual(recs.count(), 4)
        descs = {r.indicator.description for r in recs}
        self.assertIn('Publications in high-impact journals target (10)', descs)
        self.assertIn('External research funding target (50)', descs)
        self.assertIn('Patents filed target (4)', descs)
        self.assertIn('Innovation workshops target (12)', descs)

    def test_word_ovcrdes_research_block_four_kpis_merged_parents(self):
        """Two PAPs under one R&I outcome; continuation rows have empty merged Outcome/Strategy/PAP."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_ovcrdes_ri_four_kpis_merged_blanks_docx(path)
            n = ingest_word_monitor(path, office_name='OVCRDESBlock', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 4)
        recs = PerformanceRecord.objects.filter(
            quarter=1, year=2024, indicator__pap__office__name='OVCRDESBlock'
        ).select_related('indicator', 'indicator__pap')
        self.assertEqual(recs.count(), 4)
        pap_names = {r.indicator.pap.name for r in recs}
        self.assertEqual(
            pap_names,
            {
                'Young LIFTers Program for targeted students',
                'Fintech Applications Development Program for BatStateU Lipa',
            },
        )
        descs = {r.indicator.description for r in recs}
        self.assertTrue(any('students participated' in d.lower() for d in descs))
        self.assertTrue(any('writeshop' in d.lower() and 'hunt' in d.lower() for d in descs))
        self.assertTrue(any('capacity-building' in d.lower() for d in descs))
        self.assertTrue(any('concept proposal' in d.lower() and 'fintech' in d.lower() for d in descs))

    def test_word_ovcrdes_four_kpis_empty_indicator_continuation_not_skipped(self):
        """Continuation rows with pillar match but blank indicator cell must still ingest."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_ovcrdes_ri_four_kpis_empty_indicator_continuation_docx(path)
            n = ingest_word_monitor(path, office_name='OVCRDESEmptyInd', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 4)
        recs = PerformanceRecord.objects.filter(
            quarter=1, year=2024, indicator__pap__office__name='OVCRDESEmptyInd'
        ).select_related('indicator', 'indicator__pap')
        self.assertEqual(recs.count(), 4)

    def test_word_ovcrdes_four_kpis_empty_indicator_short_ac_still_ingests(self):
        """Continuation with very short accomplishment must not be dropped as non-KPI / banner."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_ovcrdes_ri_four_kpis_empty_indicator_short_narrative_docx(path)
            n = ingest_word_monitor(path, office_name='OVCRDESShortAc', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 4)
        self.assertEqual(
            PerformanceRecord.objects.filter(
                quarter=1, year=2024, indicator__pap__office__name='OVCRDESShortAc'
            ).count(),
            4,
        )

    def test_word_duplicate_indicator_cell_two_rows_two_performance_records(self):
        """Same indicator text on consecutive rows must not collapse to one DB row per quarter."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_opmm_duplicate_indicator_column_two_rows_docx(path)
            ingest_word_monitor(path, office_name='DupIndOffice', quarter=2, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        recs = list(
            PerformanceRecord.objects.filter(
                quarter=2, year=2026, indicator__pap__office__name='DupIndOffice'
            ).select_related('indicator')
        )
        self.assertEqual(len(recs), 2)
        self.assertEqual(len({r.indicator_id for r in recs}), 2)
        texts = {r.indicator.description for r in recs}
        self.assertTrue(any('Shared KPI title with target (99)' in t for t in texts))
        self.assertTrue(any('· row' in t for t in texts))

    def test_word_user_layout_banner_above_headers_two_kpis_both_ingested(self):
        """User-reported layout: pillar banner above column headers + 2 KPIs (second has merged Outcome).

        Mirrors the OVCRDES Q1 2024 monitor where ``RESEARCH AND INNOVATION: …`` spans the
        first row, column headers ``(2) Outcome … (8) Remarks`` follow, then two KPI rows
        where the second row vertically merges the Outcome cell with the first row.
        Both KPIs MUST be stored as ``PerformanceRecord`` rows.
        """
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            doc = Document()
            doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
            table = doc.add_table(rows=4, cols=7)
            banner = 'RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development'
            for ci in range(7):
                table.rows[0].cells[ci].text = banner
            for ci in range(1, 7):
                table.rows[0].cells[0].merge(table.rows[0].cells[ci])
            hdr = [
                '(2) Outcome\n(to be lifted directly from the Strategic Plan)',
                '(3) Strategy/ies\n(to be lifted directly from the Strategic Plan)',
                '(4) Action Steps/ Program, Activities, Projects (PAPs)',
                '(5) Performance Indicator/s\n(as stated in the Operational Plan)',
                '(6) Actual Accomplishments',
                '(7) Variance',
                '(8) Remarks',
            ]
            for ci, h in enumerate(hdr):
                table.rows[1].cells[ci].text = h
            r2 = table.rows[2].cells
            r2[0].text = 'University Research and Innovation Ecosystem Advanced'
            r2[1].text = 'Promote cross-disciplinary and interdisciplinary research collaborations'
            r2[2].text = 'Young Leadership and Innovation for Future Thinkers (Young LIFTERS) Program'
            r2[3].text = 'Number of students participated in the program'
            r2[4].text = (
                'One (1) implemented monitoring of research projects of YLP mentees '
                'with sufficient narrative length for ingest.'
            )
            r2[5].text = 'MET'
            r2[6].text = 'Last March 14, 2024, the RMS conducted monitoring of YLP research project.'
            r3 = table.rows[3].cells
            r2[0].merge(r3[0])
            r3[1].text = 'Strengthen capability building of researchers in different fields'
            r3[2].text = 'Regular Writeshops Kapehan and HUNTahan Session'
            r3[3].text = 'Number of conducted writeshop/HUNTahan Session'
            r3[4].text = 'Zero (0) regular Writeshops, Kapehan and HUNTahan Session'
            r3[5].text = 'N/A'
            r3[6].text = 'The target was initially set for Q1; funds may be moved to Q2.'
            doc.save(path)
            n = ingest_word_monitor(path, office_name='UserLayoutOffice', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        recs = list(
            PerformanceRecord.objects.filter(
                quarter=1, year=2024, indicator__pap__office__name='UserLayoutOffice'
            ).select_related('indicator', 'indicator__pap')
        )
        descs = sorted(r.indicator.description for r in recs)
        self.assertEqual(len(recs), 2, f'Saved descriptions: {descs}')
        self.assertEqual(n, 2)
        self.assertTrue(any('students participated' in d.lower() for d in descs))
        self.assertTrue(any('writeshop' in d.lower() for d in descs))

    def test_word_eight_column_grid_young_lifters_and_writeshop_ingested(self):
        """Tables with a leading # column must map Outcome–Remarks correctly (not shift by one)."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            doc = Document()
            doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
            table = doc.add_table(rows=5, cols=8)
            hdr = [
                '',
                'Outcome (2)',
                'Strategy/ies (3)',
                'PAP (4)',
                'Performance Indicator/s (5)',
                'Actual Accomplishments (6)',
                'Variance (7)',
                'Remarks (8)',
            ]
            for ci, h in enumerate(hdr):
                table.rows[2].cells[ci].text = h
            r3 = table.rows[3].cells
            r3[1].text = 'University Research and Innovation Ecosystem Advanced'
            r3[2].text = 'Promote cross-disciplinary and interdisciplinary research collaborations'
            r3[3].text = 'Young Leadership and Innovation for Future Thinkers (Young LIFTERS) Program'
            r3[4].text = 'Number of students participated in the program'
            r3[5].text = (
                'One (1) implemented monitoring of research projects of YLP mentees '
                'with sufficient narrative length for ingest.'
            )
            r3[6].text = 'MET'
            r4 = table.rows[4].cells
            r4[2].text = 'Strengthen capability building of researchers in different fields'
            r4[3].text = 'Regular Writeshops Kapehan and HUNTahan Session'
            r4[4].text = 'Number of conducted writeshop/HUNTahan Session'
            r4[5].text = 'Zero (0) regular Writeshops, Kapehan and HUNTahan Session'
            r4[6].text = 'N/A'
            doc.save(path)
            n = ingest_word_monitor(path, office_name='EightColGrid', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 2)
        descs = {
            r.indicator.description
            for r in PerformanceRecord.objects.filter(
                quarter=1, year=2024, indicator__pap__office__name='EightColGrid'
            )
        }
        self.assertTrue(any('students participated' in d.lower() for d in descs))
        self.assertTrue(any('writeshop' in d.lower() for d in descs))

    def test_word_ri_young_lifters_and_writeshop_both_ingested(self):
        """R&I banner row + Young LIFTERS KPI + Writeshop KPI (user report: first KPI was missing)."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_ri_banner_young_lifters_writeshop_docx(path)
            n = ingest_word_monitor(path, office_name='RiLiftersWriteshop', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 2)
        recs = list(
            PerformanceRecord.objects.filter(
                quarter=1, year=2024, indicator__pap__office__name='RiLiftersWriteshop'
            ).select_related('indicator')
        )
        self.assertEqual(len(recs), 2)
        descs = {r.indicator.description for r in recs}
        self.assertTrue(any('students participated' in d.lower() for d in descs))
        self.assertTrue(any('writeshop' in d.lower() for d in descs))

    def test_word_ri_first_kpi_outcome_filled_with_pillar_real_indicator_kept(self):
        """Outcome cell shows merged pillar text but Strategy/PAP/Indicator are real KPI columns."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            doc = Document()
            doc.add_paragraph('OPERATIONAL PLAN MONITORING MATRIX For 1st Quarter 2024')
            table = doc.add_table(rows=5, cols=7)
            hdr = [
                'Outcome (2)',
                'Strategy/ies (3)',
                'PAP (4)',
                'Performance Indicator/s (5)',
                'Actual Accomplishments (6)',
                'Variance (7)',
                'Remarks (8)',
            ]
            for ci, h in enumerate(hdr):
                table.rows[2].cells[ci].text = h
            pillar = (
                'RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development'
            )
            table.rows[3].cells[0].text = pillar
            r4 = table.rows[4].cells
            r4[0].text = pillar
            r4[1].text = 'Promote cross-disciplinary and interdisciplinary research collaborations'
            r4[2].text = 'Young Leadership and Innovation for Future Thinkers (Young LIFTERS) Program'
            r4[3].text = 'Number of students participated in the program'
            r4[4].text = (
                'One (1) implemented monitoring of research projects of YLP mentees '
                'with sufficient narrative length for ingest.'
            )
            r4[5].text = 'MET'
            doc.save(path)
            n = ingest_word_monitor(path, office_name='RiPillarOutcomeFill', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertGreaterEqual(n, 1)
        descs = {
            r.indicator.description
            for r in PerformanceRecord.objects.filter(
                quarter=1, year=2024, indicator__pap__office__name='RiPillarOutcomeFill'
            )
        }
        self.assertTrue(any('students participated' in d.lower() for d in descs))

    def test_word_ri_first_kpi_pillar_cloned_columns_still_ingested(self):
        """Pillar title repeated in Outcome–Indicator cells with real accomplishment must ingest."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_ri_first_kpi_pillar_cloned_in_hierarchy_docx(path)
            n = ingest_word_monitor(path, office_name='RiPillarCloneCols', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 2)
        descs = {
            r.indicator.description
            for r in PerformanceRecord.objects.filter(
                quarter=1, year=2024, indicator__pap__office__name='RiPillarCloneCols'
            )
        }
        self.assertTrue(any('writeshop' in d.lower() for d in descs))
        self.assertTrue(
            any('monitoring' in d.lower() or 'ylp' in d.lower() or 'row' in d.lower() for d in descs)
            or any('students' in d.lower() for d in descs)
        )

    def test_word_ri_first_kpi_duplicate_outcome_strategy_still_ingested(self):
        """Merged Outcome/Strategy duplicate text must not be classified as a pillar-only banner row."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_ri_first_kpi_duplicate_outcome_strategy_docx(path)
            n = ingest_word_monitor(path, office_name='RiDupOutcomeStrat', quarter=1, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 2)
        descs = {
            r.indicator.description
            for r in PerformanceRecord.objects.filter(
                quarter=1, year=2024, indicator__pap__office__name='RiDupOutcomeStrat'
            )
        }
        self.assertTrue(any('students participated' in d.lower() for d in descs))

    def test_word_social_banner_row_not_ingested_real_kpis_are(self):
        """Repeated pillar title + MET must not become a KPI row; following indicators must ingest."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_social_banner_row_plus_two_kpis_docx(path)
            n = ingest_word_monitor(path, office_name='SocialBannerOffice', quarter=1, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertEqual(n, 2)
        recs = list(
            PerformanceRecord.objects.filter(
                quarter=1, year=2026, indicator__pap__office__name='SocialBannerOffice'
            ).select_related('indicator')
        )
        self.assertEqual(len(recs), 2)
        for r in recs:
            self.assertFalse(record_is_pillar_banner_only(r))
            self.assertIn('Number of', r.indicator.description)

    def test_word_four_rows_identical_indicator_cell_four_performance_records(self):
        """Merged-cell repeat of the same indicator label across four rows must yield four DB rows."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_word_opmm_duplicate_indicator_cell_four_rows_docx(path)
            ingest_word_monitor(path, office_name='DupFourOffice', quarter=2, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        recs = list(
            PerformanceRecord.objects.filter(
                quarter=2, year=2026, indicator__pap__office__name='DupFourOffice'
            ).select_related('indicator')
        )
        self.assertEqual(len(recs), 4)
        self.assertEqual(len({r.indicator_id for r in recs}), 4)


def _build_lpc_wide_monitor_xlsx(path):
    """Wide FYDP / LPC layout like institutional OPMM Excel (annual target + MET column)."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.cell(row=1, column=1, value='DEVELOPMENT AREA: ACADEMIC LEADERSHIP — sample pillar')
    headers = [
        'Outcome',
        'Strategy (Based on FYDP):',
        'Program/Activity/Project (Based on FYDP):',
        'Sub PAP (Based on FYDP):',
        'Performance Indicators (based on the Cascaded Template):',
        'Performance Indicators (based on the submitted Operational Plan):',
        'Concerned Office / Campus:',
        'Q1 Target (Based on Annual Operational Plan):',
        'Q2 Target (Based on Annual Operational Plan):',
        'Q3 Target (Based on Annual Operational Plan):',
        'Q4 Target (Based on Annual Operational Plan):',
        'Annual Quantifiable Target (Based on Annual Operational Plan):',
        'Accomplishment to Date (Sum of Quarterly Accomplishments):',
        'Variance to Date (Quarterly vs Targets):',
        'Status of Accomplishment (MET or UNMET):',
    ]
    for i, h in enumerate(headers):
        ws.cell(row=2, column=i + 1, value=h)
    # Data row 3 — full hierarchy
    ws.cell(row=3, column=1, value='Outcome 1. Student services')
    ws.cell(row=3, column=2, value='Strategy 1. Improve infrastructure')
    ws.cell(row=3, column=3, value='Provision and rehabilitation of facilities')
    ws.cell(row=3, column=6, value='Percentage of renovation completed for Building A')
    ws.cell(row=3, column=7, value='LPC Office Alpha')
    ws.cell(row=3, column=12, value='25')
    ws.cell(row=3, column=13, value='33.54')
    ws.cell(row=3, column=15, value='MET')
    # Row 4 — merged simulation: empty outcome/strategy/program, same block
    ws.cell(row=4, column=6, value='Percentage Landscaping Phase II')
    ws.cell(row=4, column=7, value='LPC Office Alpha')
    ws.cell(row=4, column=12, value='50%')
    ws.cell(row=4, column=13, value='0%')
    ws.cell(row=4, column=15, value='UNMET')
    wb.save(path)
    wb.close()


class IngestExcelLpcWideTests(TestCase):
    def test_ingest_lpc_wide_splits_concerned_offices(self):
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            wb = Workbook()
            ws = wb.active
            ws.cell(row=1, column=1, value='DEVELOPMENT AREA: RESEARCH AND INNOVATION')
            hdrs = [
                'Outcome',
                'Strategy (Based on FYDP):',
                'Program/Activity/Project (Based on FYDP):',
                'Sub PAP (Based on FYDP):',
                'Performance Indicators (based on the Cascaded Template):',
                'Performance Indicators (based on the submitted Operational Plan):',
                'Concerned Office / Campus:',
                'Q1 Target',
                'Q2 Target',
                'Q3 Target',
                'Q4 Target',
                'Annual Quantifiable Target (Based on Annual Operational Plan):',
                'Accomplishment to Date (Sum of Quarterly Accomplishments):',
                'Variance',
                'Status of Accomplishment (MET or UNMET):',
            ]
            for i, h in enumerate(hdrs):
                ws.cell(row=2, column=i + 1, value=h)
            ws.cell(row=3, column=1, value='Outcome A')
            ws.cell(row=3, column=2, value='Strategy A')
            ws.cell(row=3, column=3, value='Program A')
            ws.cell(row=3, column=6, value='Indicator text one')
            ws.cell(
                row=3,
                column=7,
                value='Chancellor, OVCAF, OVCAA',
            )
            ws.cell(row=3, column=12, value='10')
            ws.cell(row=3, column=13, value='10')
            ws.cell(row=3, column=15, value='MET')
            wb.save(path)
            wb.close()
            n = ingest_excel_monitor(
                path,
                office_name='UploadDefault',
                quarter=4,
                year=2025,
            )
        finally:
            if os.path.isfile(path):
                os.unlink(path)
        self.assertEqual(n, 3)
        self.assertTrue(Office.objects.filter(name='OVCAF').exists())
        self.assertTrue(Office.objects.filter(name='OVCAA').exists())
        self.assertTrue(Office.objects.filter(name='Chancellor').exists())
        self.assertEqual(
            PerformanceRecord.objects.filter(year=2025, quarter=4).count(),
            3,
        )

    def test_ingest_lpc_wide_reads_annual_accomplishment_and_status(self):
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_lpc_wide_monitor_xlsx(path)
            n = ingest_excel_monitor(
                path,
                office_name='WideOffice',
                quarter=4,
                year=2025,
                extra_hint='LPC-OPMM-Q4-2025.xlsx',
            )
        finally:
            if os.path.isfile(path):
                os.unlink(path)
        self.assertEqual(n, 2)
        met = PerformanceRecord.objects.filter(
            year=2025, quarter=4, explicit_status='MET'
        ).count()
        unmet = PerformanceRecord.objects.filter(
            year=2025, quarter=4, explicit_status='UNMET'
        ).count()
        self.assertEqual(met, 1)
        self.assertEqual(unmet, 1)
        r_met = PerformanceRecord.objects.get(explicit_status='MET')
        self.assertEqual(r_met.target_value, 25.0)
        self.assertEqual(r_met.actual_value, 33.54)
        self.assertEqual(r_met.status, 'MET')

    def test_ingest_opmm_actual_accomplishments_and_variance_without_status_col(self):
        """OPMM-style: Actual Accomplishments + Variance (MET / +N); no Status column."""
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            wb = Workbook()
            ws = wb.active
            ws.cell(row=1, column=1, value='DEVELOPMENT AREA: SOCIAL RESPONSIBILITY — sample')
            hdrs = [
                'Outcome',
                'Strategy (Based on FYDP):',
                'Program/Activity/Project (Based on FYDP):',
                'Sub PAP (Based on FYDP):',
                'Performance Indicators (based on the Cascaded Template):',
                'Performance Indicators (based on the submitted Operational Plan):',
                'Concerned Office / Campus:',
                'Q1 Target (Based on Annual Operational Plan):',
                'Q2 Target (Based on Annual Operational Plan):',
                'Q3 Target (Based on Annual Operational Plan):',
                'Q4 Target (Based on Annual Operational Plan):',
                'Annual Quantifiable Target (Based on Annual Operational Plan):',
                'Actual Accomplishments',
                'Variance',
            ]
            for i, h in enumerate(hdrs):
                ws.cell(row=2, column=i + 1, value=h)
            ws.cell(row=3, column=1, value='Outcome one')
            ws.cell(row=3, column=2, value='Strategy one')
            ws.cell(row=3, column=3, value='Program one')
            ws.cell(row=3, column=6, value='Gender PPAs (2)')
            ws.cell(row=3, column=12, value='2')
            ws.cell(row=3, column=13, value='Two PPAs completed')
            ws.cell(row=3, column=14, value='MET')
            ws.cell(row=4, column=6, value='Radio faculty (4)')
            ws.cell(row=4, column=12, value='4')
            ws.cell(row=4, column=13, value='Five episodes')
            ws.cell(row=4, column=14, value='+3')
            wb.save(path)
            wb.close()
            n = ingest_excel_monitor(path, office_name='OpmmStyle', quarter=4, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)
        self.assertEqual(n, 2)
        self.assertEqual(
            PerformanceRecord.objects.filter(
                year=2024, quarter=4, explicit_status='MET'
            ).count(),
            2,
        )

    def test_ingest_opmm_monitoring_matrix_sparse_parents_no_annual_column(self):
        """BatStateU-style Q4 matrix: (n) headers, Actual + Variance, blank parent cells on continuation rows."""
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            wb = Workbook()
            ws = wb.active
            ws.cell(
                row=1,
                column=1,
                value='RESEARCH AND INNOVATION: Engineering Innovative Solutions for Sustainable Development.',
            )
            hdrs = [
                '(2) Outcome',
                '(3) Strategy/ies',
                '(4) Action Steps/ Program, Activities, Projects (PAPs):',
                '(5) Performance Indicator/s',
                '(6) Actual Accomplishments',
                '(7) Variance',
                '(8) Remarks',
            ]
            for i, h in enumerate(hdrs):
                ws.cell(row=2, column=i + 1, value=h)
            ws.cell(row=3, column=1, value='University Research and Innovation Ecosystem Advanced')
            ws.cell(row=3, column=2, value='Establish enabling policies and environment')
            ws.cell(row=3, column=3, value='Publication Impact and Productivity Incentives (π²)')
            ws.cell(row=3, column=4, value='Number of Publications in high-impact research journals')
            ws.cell(row=3, column=5, value='Twelve (12) publications')
            ws.cell(row=3, column=6, value='MET')
            ws.cell(row=3, column=7, value='See annex A')
            # Continuation: same PAP, blank outcome/strategy/pap columns
            ws.cell(row=4, column=4, value='Number of submitted proposal w/ Budget higher than 150,000.00')
            ws.cell(row=4, column=5, value='Four (4) submitted')
            ws.cell(row=4, column=6, value='-1')
            ws.cell(row=4, column=7, value='')
            ws.cell(row=5, column=4, value='Number of approved proposal w/ budget higher than 150,000.00')
            ws.cell(row=5, column=5, value='Three (3) approved')
            ws.cell(row=5, column=6, value='UNMET')
            ws.cell(row=5, column=7, value='Target was 4')
            wb.save(path)
            wb.close()
            n = ingest_excel_monitor(path, office_name='OvcrmSample', quarter=4, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)
        self.assertEqual(n, 3)
        self.assertEqual(
            PerformanceRecord.objects.filter(
                indicator__pap__office__name='OvcrmSample', year=2024, quarter=4
            ).count(),
            3,
        )
        self.assertEqual(
            PerformanceRecord.objects.filter(
                indicator__pap__office__name='OvcrmSample',
                year=2024,
                quarter=4,
                explicit_status='MET',
            ).count(),
            1,
        )
        self.assertEqual(
            PerformanceRecord.objects.filter(
                indicator__pap__office__name='OvcrmSample',
                year=2024,
                quarter=4,
                explicit_status='UNMET',
            ).count(),
            2,
        )
        pap_names = set(
            Indicator.objects.filter(pap__office__name='OvcrmSample').values_list(
                'pap__name', flat=True
            )
        )
        self.assertIn('Publication Impact and Productivity Incentives (π²)', pap_names)

    def test_ingest_lpc_wide_sums_rows_across_worksheets(self):
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            wb = Workbook()
            ws1 = wb.active
            ws1.title = 'Social'
            ws1.cell(row=1, column=1, value='DEVELOPMENT AREA: SOCIAL RESPONSIBILITY')
            h1 = [
                'Outcome',
                'Strategy (Based on FYDP):',
                'Program/Activity/Project (Based on FYDP):',
                'Sub PAP (Based on FYDP):',
                'Performance Indicators (based on the Cascaded Template):',
                'Performance Indicators (based on the submitted Operational Plan):',
                'Concerned Office / Campus:',
                'Q1 Target (Based on Annual Operational Plan):',
                'Q2 Target (Based on Annual Operational Plan):',
                'Q3 Target (Based on Annual Operational Plan):',
                'Q4 Target (Based on Annual Operational Plan):',
                'Annual Quantifiable Target (Based on Annual Operational Plan):',
                'Accomplishment to Date (Sum of Quarterly Accomplishments):',
                'Variance',
            ]
            for i, h in enumerate(h1):
                ws1.cell(row=2, column=i + 1, value=h)
            ws1.cell(row=3, column=1, value='O1')
            ws1.cell(row=3, column=2, value='S1')
            ws1.cell(row=3, column=3, value='P1')
            ws1.cell(row=3, column=6, value='Ind sheet1 (5)')
            ws1.cell(row=3, column=12, value='5')
            ws1.cell(row=3, column=13, value='5')
            ws1.cell(row=3, column=14, value='MET')
            ws2 = wb.create_sheet('Research')
            ws2.cell(row=1, column=1, value='DEVELOPMENT AREA: RESEARCH AND INNOVATION')
            for i, h in enumerate(h1):
                ws2.cell(row=2, column=i + 1, value=h)
            ws2.cell(row=3, column=1, value='RO')
            ws2.cell(row=3, column=2, value='RS')
            ws2.cell(row=3, column=3, value='RP')
            ws2.cell(row=3, column=6, value='Ind sheet2 (10)')
            ws2.cell(row=3, column=12, value='10')
            ws2.cell(row=3, column=13, value='12')
            ws2.cell(row=3, column=14, value='MET')
            wb.save(path)
            wb.close()
            n = ingest_excel_monitor(path, office_name='TwoSheetOffice', quarter=4, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)
        self.assertEqual(n, 2)
        self.assertEqual(PerformanceRecord.objects.filter(year=2026, quarter=4).count(), 2)

    def test_ingest_lpc_wide_merged_indicator_column_does_not_collapse_rows(self):
        """Merged indicator cells leave blanks on continuation rows; each KPI must be its own record."""
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            wb = Workbook()
            ws = wb.active
            ws.cell(row=1, column=1, value='DEVELOPMENT AREA: RESEARCH AND INNOVATION')
            h1 = [
                'Outcome',
                'Strategy (Based on FYDP):',
                'Program/Activity/Project (Based on FYDP):',
                'Sub PAP (Based on FYDP):',
                'Performance Indicators (based on the Cascaded Template):',
                'Performance Indicators (based on the submitted Operational Plan):',
                'Concerned Office / Campus:',
                'Q1 Target (Based on Annual Operational Plan):',
                'Q2 Target (Based on Annual Operational Plan):',
                'Q3 Target (Based on Annual Operational Plan):',
                'Q4 Target (Based on Annual Operational Plan):',
                'Annual Quantifiable Target (Based on Annual Operational Plan):',
                'Accomplishment to Date (Sum of Quarterly Accomplishments):',
                'Variance',
                'Status of Accomplishment (MET or UNMET):',
            ]
            for i, h in enumerate(h1):
                ws.cell(row=2, column=i + 1, value=h)
            for r, ind_cell, ann, acmp, st in (
                (3, 'Shared merged indicator label', '5', 'First row acmp', 'MET'),
                (4, '', '8', 'Second row acmp', 'UNMET'),
                (5, '', '9', 'Third row acmp', 'MET'),
            ):
                ws.cell(row=r, column=1, value='O-merge')
                ws.cell(row=r, column=2, value='S-merge')
                ws.cell(row=r, column=3, value='P-merge')
                ws.cell(row=r, column=6, value=ind_cell)
                ws.cell(row=r, column=12, value=ann)
                ws.cell(row=r, column=13, value=acmp)
                ws.cell(row=r, column=15, value=st)
            wb.save(path)
            wb.close()
            n = ingest_excel_monitor(path, office_name='MergeIndOffice', quarter=4, year=2024)
        finally:
            if os.path.isfile(path):
                os.unlink(path)
        self.assertEqual(n, 3)
        self.assertEqual(
            PerformanceRecord.objects.filter(
                year=2024, quarter=4, indicator__pap__office__name='MergeIndOffice'
            ).count(),
            3,
        )
        descs = set(
            PerformanceRecord.objects.filter(
                year=2024, quarter=4, indicator__pap__office__name='MergeIndOffice'
            ).values_list('indicator__description', flat=True)
        )
        self.assertEqual(len(descs), 3)
        self.assertIn('Shared merged indicator label', descs)


class IngestExcelMonitorTests(TestCase):
    def test_ingest_xlsx_creates_performance_row(self):
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_min_monitor_xlsx(path)
            ingest_excel_monitor(path, office_name='ExcelOffice', quarter=2, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        self.assertTrue(Office.objects.filter(name='ExcelOffice').exists())
        rec = PerformanceRecord.objects.filter(quarter=2, year=2026).first()
        self.assertIsNotNone(rec)
        self.assertEqual(rec.target_value, 100.0)
        self.assertEqual(rec.actual_value, 95.0)
        self.assertEqual(rec.status, 'UNMET')

    def test_opmm_matrix_two_strategies_multiple_indicators_per_pap(self):
        """One outcome, two strategies, two PAPs, multiple indicators; strategy change clears PAP carry."""
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            wb = Workbook()
            ws = wb.active
            hdrs = [
                '(2) Outcome',
                '(3) Strategy/ies',
                '(4) Action Steps/ Program, Activities, Projects (PAPs):',
                '(5) Performance Indicator/s',
                '(6) Actual Accomplishments',
                '(7) Variance',
                '(8) Remarks',
            ]
            for i, h in enumerate(hdrs):
                ws.cell(row=1, column=i + 1, value=h)
            ws.cell(row=2, column=1, value='Single Outcome X')
            ws.cell(row=2, column=2, value='Strategy One')
            ws.cell(row=2, column=3, value='PAP Alpha')
            ws.cell(row=2, column=4, value='Indicator A1 under PAP Alpha')
            ws.cell(row=2, column=5, value='Done (1)')
            ws.cell(row=2, column=6, value='MET')
            ws.cell(row=3, column=4, value='Indicator A2 under PAP Alpha')
            ws.cell(row=3, column=5, value='Done (2)')
            ws.cell(row=3, column=6, value='UNMET')
            ws.cell(row=4, column=2, value='Strategy Two')
            ws.cell(row=4, column=3, value='PAP Beta')
            ws.cell(row=4, column=4, value='Indicator B1 under PAP Beta')
            ws.cell(row=4, column=5, value='Done (3)')
            ws.cell(row=4, column=6, value='MET')
            ws.cell(row=5, column=4, value='Indicator B2 under PAP Beta')
            ws.cell(row=5, column=5, value='Done (4)')
            ws.cell(row=5, column=6, value='MET')
            wb.save(path)
            wb.close()
            n = ingest_excel_monitor(path, office_name='MultiStratOffice', quarter=2, year=2025)
        finally:
            if os.path.isfile(path):
                os.unlink(path)
        self.assertEqual(n, 4)
        p_alpha = Indicator.objects.filter(
            pap__office__name='MultiStratOffice', pap__name='PAP Alpha'
        )
        p_beta = Indicator.objects.filter(
            pap__office__name='MultiStratOffice', pap__name='PAP Beta'
        )
        self.assertEqual(p_alpha.count(), 2)
        self.assertEqual(p_beta.count(), 2)
        self.assertEqual(
            set(p_alpha.values_list('description', flat=True)),
            {
                'Indicator A1 under PAP Alpha',
                'Indicator A2 under PAP Alpha',
            },
        )
        self.assertEqual(
            set(p_beta.values_list('description', flat=True)),
            {
                'Indicator B1 under PAP Beta',
                'Indicator B2 under PAP Beta',
            },
        )
        self.assertEqual(
            PerformanceRecord.objects.filter(
                indicator__in=p_alpha, explicit_status='MET'
            ).count(),
            1,
        )
        self.assertEqual(
            PerformanceRecord.objects.filter(
                indicator__in=p_alpha, explicit_status='UNMET'
            ).count(),
            1,
        )

    def test_ingest_xlsx_resolves_period_from_filename_hint(self):
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            _build_min_monitor_xlsx(path)
            ingest_excel_monitor(
                path,
                office_name='ExcelHint',
                quarter=None,
                year=None,
                extra_hint='LPC-OPMM-Q4-2025.xlsx',
            )
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        rec = PerformanceRecord.objects.filter(
            indicator__pap__office__name='ExcelHint',
        ).first()
        self.assertIsNotNone(rec)
        self.assertEqual(rec.quarter, 4)
        self.assertEqual(rec.year, 2025)

    def test_ingest_xlsx_office_column_f(self):
        from openpyxl import Workbook

        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            path = tmp.name
        try:
            wb = Workbook()
            ws = wb.active
            for ri in range(3):
                for ci in range(5):
                    v = f'header_{ri}_{ci}'
                    if ri == 0 and ci == 0:
                        v = 'Status Report Q1 2026'
                    ws.cell(row=ri + 1, column=ci + 1, value=v)
            ws.cell(row=4, column=1, value='Academic Leadership and excellence')
            ws.cell(row=4, column=2, value='Strategy Alpha')
            ws.cell(row=4, column=3, value='PAP Sample')
            ws.cell(row=4, column=4, value='Indicator target (50)')
            ws.cell(row=4, column=5, value='Accomplishment (50)')
            ws.cell(row=4, column=6, value='PDO From Column F')
            wb.save(path)
            ingest_excel_monitor(path, office_name='FallbackOffice', quarter=1, year=2026)
        finally:
            if os.path.isfile(path):
                os.unlink(path)

        rec = PerformanceRecord.objects.filter(quarter=1, year=2026).first()
        self.assertIsNotNone(rec)
        self.assertEqual(rec.indicator.pap.office.name, 'PDO From Column F')


@override_settings(DEBUG=True)
class ClearDataRouteTests(TestCase):
    """DEBUG=True so clears succeed regardless of host DJANGO_DEBUG."""

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            username='adm',
            password='pass12345',
            is_superuser=True,
            is_staff=True,
        )

    def test_clear_all_data_rejects_get(self):
        self.client.login(username='adm', password='pass12345')
        url = reverse('clear_all_data')
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 405)

    def test_clear_all_data_post_clears_models(self):
        Office.objects.create(name='O1')
        self.client.login(username='adm', password='pass12345')
        self.assertEqual(PerformanceRecord.objects.count(), 0)
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            ingest_word_monitor(tmp.name, office_name='O1', quarter=1, year=2026)
            os.unlink(tmp.name)
        self.assertGreater(PerformanceRecord.objects.count(), 0)

        resp = self.client.post(reverse('clear_all_data'))
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(PerformanceRecord.objects.count(), 0)
        self.assertEqual(Indicator.objects.count(), 0)
        self.assertEqual(StrategicLevel.objects.count(), 0)

    @override_settings(DEBUG=False, SOPM_ENABLE_FULL_DATABASE_CLEAR=False)
    def test_clear_all_blocked_in_production_without_flag(self):
        self.client.login(username='adm', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            ingest_word_monitor(tmp.name, office_name='O2', quarter=1, year=2026)
            os.unlink(tmp.name)
        before = PerformanceRecord.objects.count()
        self.assertGreater(before, 0)
        self.client.post(reverse('clear_all_data'))
        self.assertEqual(PerformanceRecord.objects.count(), before)

    @override_settings(DEBUG=False, SOPM_ENABLE_FULL_DATABASE_CLEAR=True)
    def test_clear_all_allowed_when_env_flag_enabled(self):
        self.client.login(username='adm', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            ingest_word_monitor(tmp.name, office_name='O3', quarter=1, year=2026)
            os.unlink(tmp.name)
        self.client.post(reverse('clear_all_data'))
        self.assertEqual(PerformanceRecord.objects.count(), 0)

    def test_clear_office_data_rejects_get(self):
        Office.objects.create(name='Staff Office Full Name')
        staff = User.objects.create_user(
            username='stf',
            password='pass12345',
            first_name='Staff Office Full Name',
        )
        self.client.login(username='stf', password='pass12345')
        self.assertEqual(self.client.get(reverse('clear_office_data')).status_code, 405)


class OfficeChartPayloadTests(TestCase):
    def test_orders_offices_by_met_then_fewer_unmet(self):
        out = StrategicLevel.objects.create(name='Academic Leadership', level_type='OUTCOME')
        st = StrategicLevel.objects.create(name='Strategy', level_type='STRATEGY', parent=out)
        o1 = Office.objects.create(name='Alpha Office')
        o2 = Office.objects.create(name='Beta Office')
        pap1 = StrategicLevel.objects.create(name='PAP1', level_type='PAP', parent=st, office=o1)
        pap2 = StrategicLevel.objects.create(name='PAP2', level_type='PAP', parent=st, office=o2)
        i1 = Indicator.objects.create(pap=pap1, description='T (10)')
        i2 = Indicator.objects.create(pap=pap1, description='T (10)')
        i3 = Indicator.objects.create(pap=pap2, description='T (10)')
        PerformanceRecord.objects.create(
            indicator=i1, quarter=1, year=2026, target_value=10, actual_value=10
        )
        PerformanceRecord.objects.create(
            indicator=i2, quarter=1, year=2026, target_value=10, actual_value=0
        )
        PerformanceRecord.objects.create(
            indicator=i3, quarter=1, year=2026, target_value=10, actual_value=10
        )
        qs = PerformanceRecord.objects.select_related('indicator__pap__office').filter(
            year=2026, quarter=1
        )
        payload = _build_office_chart_payload(list(qs))
        self.assertEqual(len(payload['labels']), 2)
        self.assertEqual(payload['labels'][0], 'Beta Office')
        self.assertEqual(payload['met'], [1, 1])
        self.assertEqual(payload['unmet'], [0, 1])

    def test_single_office_with_data_returns_one_label(self):
        out = StrategicLevel.objects.create(name='Academic Leadership', level_type='OUTCOME')
        st = StrategicLevel.objects.create(name='Strategy', level_type='STRATEGY', parent=out)
        o1 = Office.objects.create(name='Solo Office')
        pap1 = StrategicLevel.objects.create(name='PAP1', level_type='PAP', parent=st, office=o1)
        i1 = Indicator.objects.create(pap=pap1, description='T (10)')
        PerformanceRecord.objects.create(
            indicator=i1, quarter=1, year=2026, target_value=10, actual_value=10
        )
        qs = PerformanceRecord.objects.select_related('indicator__pap__office').filter(
            year=2026, quarter=1
        )
        payload = _build_office_chart_payload(list(qs))
        self.assertEqual(payload['labels'], ['Solo Office'])
        self.assertEqual(payload['totals'], [1])


class SeedSampleDataCommandTests(TestCase):
    def test_seed_reset_creates_demo_offices_and_rows(self):
        call_command('seed_sample_data', '--reset', verbosity=0)
        self.assertEqual(Office.objects.filter(name__startswith='Demo —').count(), 3)
        self.assertGreater(PerformanceRecord.objects.count(), 0)

    def test_seed_bulk_creates_many_rows(self):
        call_command(
            'seed_sample_data',
            '--reset',
            '--bulk-indicators',
            '100',
            '--bulk-quarters',
            '4',
            verbosity=0,
        )
        self.assertGreaterEqual(Indicator.objects.count(), 100)
        self.assertGreaterEqual(PerformanceRecord.objects.count(), 100 * 4)

    def test_seed_stress_flag_matches_bulk_100_by_4(self):
        call_command('seed_sample_data', '--reset', '--stress', verbosity=0)
        self.assertGreaterEqual(Indicator.objects.count(), 100)
        self.assertGreaterEqual(PerformanceRecord.objects.count(), 100 * 4)

    def test_seed_balanced_per_area_even_met_split(self):
        call_command(
            'seed_sample_data',
            '--reset',
            '--balanced-per-area',
            '2',
            '--balanced-met-pct',
            '80',
            verbosity=0,
        )
        even = PerformanceRecord.objects.filter(
            raw_actual_text='even-demo', year=2026, quarter=1
        )
        self.assertEqual(even.count(), 12)
        met = sum(1 for r in even if r.status == 'MET')
        self.assertEqual(met, 11)

    def test_seed_balanced_only_skips_curated_seed(self):
        call_command(
            'seed_sample_data',
            '--reset',
            '--balanced-only',
            '--balanced-per-area',
            '5',
            verbosity=0,
        )
        self.assertEqual(Indicator.objects.count(), 30)
        self.assertEqual(
            PerformanceRecord.objects.filter(raw_actual_text='even-demo').count(),
            30,
        )

    def test_seed_balanced_only_rejects_bulk(self):
        with self.assertRaises(CommandError):
            call_command(
                'seed_sample_data',
                '--reset',
                '--balanced-only',
                '--balanced-per-area',
                '5',
                '--bulk-indicators',
                '10',
                verbosity=0,
            )

    def test_seed_without_reset_errors_if_strategic_data_exists(self):
        StrategicLevel.objects.create(name='Existing', level_type='OUTCOME')
        with self.assertRaises(CommandError):
            call_command('seed_sample_data', verbosity=0)


class SelectiveOfficeResetTests(TestCase):
    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            username='superdel',
            password='pass12345',
            is_superuser=True,
            is_staff=True,
        )
        self.staff = User.objects.create_user(
            username='staffonly',
            password='pass12345',
            is_superuser=False,
        )

    def test_non_superuser_redirected(self):
        self.client.login(username='staffonly', password='pass12345')
        resp = self.client.get(reverse('selective_office_reset'))
        self.assertEqual(resp.status_code, 302)

    def test_superuser_get_renders(self):
        self.client.login(username='superdel', password='pass12345')
        resp = self.client.get(reverse('selective_office_reset'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Remove office data')

    def test_post_requires_delete_phrase(self):
        self.client.login(username='superdel', password='pass12345')
        out = StrategicLevel.objects.create(name='Academic Leadership', level_type='OUTCOME')
        st = StrategicLevel.objects.create(name='Strategy', level_type='STRATEGY', parent=out)
        o = Office.objects.create(name='Del Office')
        StrategicLevel.objects.create(name='PAP X', level_type='PAP', parent=st, office=o)
        oid = o.id
        resp = self.client.post(
            reverse('selective_office_reset'),
            {'office_ids': [str(oid)], 'confirm_phrase': 'remove'},
        )
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'DELETE')
        self.assertTrue(StrategicLevel.objects.filter(office=o, level_type='PAP').exists())

    def test_post_delete_removes_office_paps(self):
        self.client.login(username='superdel', password='pass12345')
        out = StrategicLevel.objects.create(name='Academic Leadership', level_type='OUTCOME')
        st = StrategicLevel.objects.create(name='Strategy', level_type='STRATEGY', parent=out)
        o = Office.objects.create(name='Zap Office')
        pap = StrategicLevel.objects.create(name='PAP Z', level_type='PAP', parent=st, office=o)
        ind = Indicator.objects.create(pap=pap, description='Ind (10)')
        PerformanceRecord.objects.create(
            indicator=ind, quarter=1, year=2026, target_value=10, actual_value=10
        )
        oid = o.id
        resp = self.client.post(
            reverse('selective_office_reset'),
            {'office_ids': [str(oid)], 'confirm_phrase': 'DELETE'},
        )
        self.assertRedirects(resp, reverse('selective_office_reset'), fetch_redirect_response=False)
        self.assertFalse(StrategicLevel.objects.filter(office=o, level_type='PAP').exists())
        self.assertEqual(PerformanceRecord.objects.count(), 0)


class UserManagementEditTests(TestCase):
    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            username='adminroot',
            password='pass12345',
            is_superuser=True,
            is_staff=True,
        )
        self.staff = User.objects.create_user(
            username='PDO',
            password='oldpw',
            first_name='Planning And Development Office',
        )
        self.office = Office.objects.create(
            name='Planning And Development Office',
            code='PDO',
        )

    def test_edit_user_get_renders_form(self):
        self.client.login(username='adminroot', password='pass12345')
        resp = self.client.get(reverse('edit_user', args=[self.staff.pk]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Edit office account')
        self.assertContains(resp, 'Planning And Development Office')

    def test_edit_user_post_updates_user_and_linked_office(self):
        self.client.login(username='adminroot', password='pass12345')
        resp = self.client.post(
            reverse('edit_user', args=[self.staff.pk]),
            {
                'full_name': 'Planning And Development Office Renamed',
                'abbreviation': 'PDOR',
                'password': 'newpw99999',
            },
        )
        self.assertRedirects(resp, reverse('user_management'), fetch_redirect_response=False)
        self.staff.refresh_from_db()
        self.assertEqual(self.staff.username, 'PDOR')
        self.assertEqual(self.staff.first_name, 'Planning And Development Office Renamed')
        self.assertTrue(self.staff.check_password('newpw99999'))
        self.office.refresh_from_db()
        self.assertEqual(self.office.name, 'Planning And Development Office Renamed')
        self.assertEqual(self.office.code, 'PDOR')


class DashboardPlanHierarchyViewTests(TestCase):
    """Office dashboard can group KPIs under Outcome → Strategy → PAP (every row visible)."""

    def setUp(self):
        self.client = Client()
        self.office = Office.objects.create(name='HierOffice')
        self.user = User.objects.create_user(
            username='hieruser',
            password='pwtest12',
            first_name='HierOffice',
        )
        out = StrategicLevel.objects.create(name='Outcome O', level_type='OUTCOME')
        strat = StrategicLevel.objects.create(name='Strategy S', level_type='STRATEGY', parent=out)
        pap = StrategicLevel.objects.create(
            name='PAP P',
            level_type='PAP',
            parent=strat,
            office=self.office,
        )
        i1 = Indicator.objects.create(pap=pap, description='IND line one')
        i2 = Indicator.objects.create(pap=pap, description='IND line two')
        PerformanceRecord.objects.create(
            indicator=i1,
            quarter=4,
            year=2024,
            target_value=1.0,
            actual_value=2.0,
            raw_actual_text='2',
        )
        PerformanceRecord.objects.create(
            indicator=i2,
            quarter=4,
            year=2024,
            target_value=2.0,
            actual_value=1.0,
            raw_actual_text='1',
        )

    def test_plan_layout_shows_full_hierarchy_and_both_rows(self):
        self.client.login(username='hieruser', password='pwtest12')
        resp = self.client.get(
            reverse('dashboard_home'),
            {'quarter': '4', 'year': '2024'},
            follow=True,
        )
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(resp.redirect_chain)
        self.assertContains(resp, 'Outcome O')
        self.assertContains(resp, 'Strategy S')
        self.assertContains(resp, 'PAP P')
        self.assertContains(resp, 'IND line one')
        self.assertContains(resp, 'IND line two')

    def test_office_user_bare_dashboard_redirects_to_all_quarters_and_all_areas(self):
        self.client.login(username='hieruser', password='pwtest12')
        resp = self.client.get(reverse('dashboard_home'), follow=False)
        self.assertEqual(resp.status_code, 302)
        url = resp.url or ''
        self.assertIn('/performance-viewer/', url)
        self.assertIn('quarter=all', url)
        self.assertIn('area=all', url)
        self.assertIn('year=2024', url)
        r2 = self.client.get(reverse('performance_viewer'), follow=False)
        self.assertEqual(r2.status_code, 200)
        self.assertContains(r2, 'All quarters')
        self.assertContains(r2, 'All development areas')

    def test_dashboard_renders_matrix_table_columns(self):
        self.client.login(username='hieruser', password='pwtest12')
        resp = self.client.get(
            reverse('dashboard_home'),
            {'quarter': '4', 'year': '2024'},
            follow=True,
        )
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(resp.redirect_chain)
        self.assertContains(resp, 'Indicator matrix')
        self.assertContains(resp, 'Actual accomplishments')
        self.assertContains(resp, 'IND line one')


class ViewSmokeTests(TestCase):
    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            username='root',
            password='pass12345',
            is_superuser=True,
            is_staff=True,
        )

    def test_root_redirects_anonymous_to_login(self):
        resp = self.client.get(reverse('root'), follow=False)
        self.assertEqual(resp.status_code, 302)
        self.assertIn(settings.LOGIN_URL, resp.url or '')

    def test_dashboard_redirects_superuser_to_performance_viewer(self):
        self.client.login(username='root', password='pass12345')
        resp = self.client.get(reverse('dashboard_home'), follow=False)
        self.assertEqual(resp.status_code, 302)
        self.assertIn('/performance-viewer/', resp.url or '')
        self.assertIn('quarter=all', resp.url or '')
        self.assertIn('area=all', resp.url or '')
        self.assertRegex(resp.url or '', r'[?&]year=\d+')

    def test_performance_viewer_renders_for_superuser(self):
        self.client.login(username='root', password='pass12345')
        resp = self.client.get(reverse('performance_viewer'))
        self.assertEqual(resp.status_code, 200)

    def test_upload_get_renders(self):
        self.client.login(username='root', password='pass12345')
        resp = self.client.get(reverse('upload_blueprint'))
        self.assertEqual(resp.status_code, 200)

    def test_upload_post_docx_redirects_success(self):
        self.client.login(username='root', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            with open(tmp.name, 'rb') as fh:
                data = fh.read()
            os.unlink(tmp.name)
        upload = SimpleUploadedFile('blueprint.docx', data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp = self.client.post(
            reverse('upload_blueprint'),
            {
                'blueprint': upload,
                'ingest_quarter': '4',
                'ingest_year': '2026',
            },
        )
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(resp.url.startswith(f'{reverse("success_page")}?n=1&rows=1'))
        self.assertIn('&ind=', resp.url)
        self.assertIn('&met=', resp.url)
        self.assertIn('&unmet=', resp.url)
        rec = PerformanceRecord.objects.filter(quarter=4, year=2026).first()
        self.assertIsNotNone(rec)

    def test_upload_staff_resolves_office_name_case_insensitive(self):
        Office.objects.create(name='Office of Research', code='RESEARCH')
        User.objects.create_user(
            username='res',
            password='pass12345',
            first_name='office of research',
        )
        self.client.login(username='res', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            with open(tmp.name, 'rb') as fh:
                data = fh.read()
            os.unlink(tmp.name)
        upload = SimpleUploadedFile(
            'blueprint.docx',
            data,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        resp = self.client.post(
            reverse('upload_blueprint'),
            {
                'blueprint': upload,
                'ingest_quarter': '4',
                'ingest_year': '2026',
            },
        )
        self.assertEqual(resp.status_code, 302)
        rec = PerformanceRecord.objects.filter(quarter=4, year=2026).first()
        self.assertIsNotNone(rec)
        self.assertEqual(rec.indicator.pap.office.name, 'Office of Research')

    def test_upload_staff_resolves_office_by_username_code_when_name_mismatches(self):
        Office.objects.create(name='Office of Something Official', code='MYOFC')
        User.objects.create_user(
            username='MYOFC',
            password='pass12345',
            first_name='Typo Full Name That Does Not Match DB',
        )
        self.client.login(username='MYOFC', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            with open(tmp.name, 'rb') as fh:
                data = fh.read()
            os.unlink(tmp.name)
        upload = SimpleUploadedFile(
            'blueprint.docx',
            data,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        resp = self.client.post(
            reverse('upload_blueprint'),
            {
                'blueprint': upload,
                'ingest_quarter': '1',
                'ingest_year': '2026',
            },
        )
        self.assertEqual(resp.status_code, 302)
        rec = PerformanceRecord.objects.filter(quarter=1, year=2026).first()
        self.assertIsNotNone(rec)
        self.assertEqual(rec.indicator.pap.office.name, 'Office of Something Official')

    def test_upload_staff_auto_creates_office_when_none_linked(self):
        User.objects.create_user(
            username='freshlogin',
            password='pass12345',
            first_name='',
        )
        self.assertFalse(Office.objects.filter(name='freshlogin').exists())
        self.client.login(username='freshlogin', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            with open(tmp.name, 'rb') as fh:
                data = fh.read()
            os.unlink(tmp.name)
        upload = SimpleUploadedFile(
            'blueprint.docx',
            data,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        resp = self.client.post(
            reverse('upload_blueprint'),
            {
                'blueprint': upload,
                'ingest_quarter': '2',
                'ingest_year': '2026',
            },
        )
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(Office.objects.filter(name='freshlogin').exists())
        rec = PerformanceRecord.objects.filter(quarter=2, year=2026).first()
        self.assertIsNotNone(rec)
        self.assertEqual(rec.indicator.pap.office.name, 'freshlogin')

    def test_upload_post_xlsx_redirects_success(self):
        self.client.login(username='root', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_xlsx(tmp.name)
            with open(tmp.name, 'rb') as fh:
                data = fh.read()
            os.unlink(tmp.name)
        upload = SimpleUploadedFile(
            'blueprint.xlsx',
            data,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        resp = self.client.post(
            reverse('upload_blueprint'),
            {
                'blueprint': upload,
                'ingest_quarter': '3',
                'ingest_year': '2026',
            },
        )
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(resp.url.startswith(f'{reverse("success_page")}?n=1&rows=1'))
        self.assertIn('&ind=', resp.url)
        self.assertIn('&met=', resp.url)
        self.assertIn('&unmet=', resp.url)
        rec = PerformanceRecord.objects.filter(quarter=3, year=2026).first()
        self.assertIsNotNone(rec)

    def test_upload_post_multiple_docx_redirects_with_count(self):
        self.client.login(username='root', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            with open(tmp.name, 'rb') as fh:
                data = fh.read()
            os.unlink(tmp.name)
        u1 = SimpleUploadedFile(
            'office_a.docx',
            data,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        u2 = SimpleUploadedFile(
            'office_b.docx',
            data,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        resp = self.client.post(
            reverse('upload_blueprint'),
            {
                'blueprint': [u1, u2],
                'ingest_quarter': '1',
                'ingest_year': '2026',
            },
        )
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(resp.url.startswith(f'{reverse("success_page")}?n=2&rows=2'))
        self.assertIn('&ind=', resp.url)
        self.assertIn('&met=', resp.url)
        self.assertIn('&unmet=', resp.url)

    def test_performance_viewer_renders_indicator_matrix_after_upload(self):
        self.client.login(username='root', password='pass12345')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.close()
            _build_min_monitor_docx(tmp.name)
            with open(tmp.name, 'rb') as fh:
                data = fh.read()
            os.unlink(tmp.name)
        upload = SimpleUploadedFile(
            'blueprint.docx',
            data,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.client.post(
            reverse('upload_blueprint'),
            {'blueprint': upload, 'ingest_quarter': '2', 'ingest_year': '2026'},
        )
        resp = self.client.get(reverse('performance_viewer') + '?year=2026&quarter=2')
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Indicator matrix')
        self.assertContains(resp, 'viewer-table')
        self.assertNotContains(resp, 'officeStatusDonut')

    def test_export_lpc_wide_excel_download(self):
        call_command('seed_sample_data', '--reset', verbosity=0)
        self.client.login(username='root', password='pass12345')
        resp = self.client.get(reverse('export_lpc_wide_excel') + '?year=2026')
        self.assertEqual(resp.status_code, 200)
        self.assertIn(
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            resp.get('Content-Type', ''),
        )
        self.assertTrue(resp.content.startswith(b'PK'))
        from openpyxl import load_workbook

        import io

        wb = load_workbook(io.BytesIO(resp.content))
        ws = wb.active
        self.assertEqual((ws.cell(row=2, column=1).value or '').strip(), 'Outcome')
        self.assertGreaterEqual(ws.max_row, 3)
        from core.services import _detect_lpc_wide_layout

        self.assertIsNotNone(_detect_lpc_wide_layout(ws))
        from core.models import PerformanceRecord

        n_db = PerformanceRecord.objects.filter(year=2026).count()
        self.assertEqual(ws.max_row, 2 + n_db)

    def test_performance_viewer_scopes_kpi_to_selected_area(self):
        call_command('seed_sample_data', '--reset', verbosity=0)
        self.client.login(username='root', password='pass12345')
        r_sust = self.client.get(
            reverse('performance_viewer') + '?year=2026&quarter=all&area=sustainability',
        )
        r_intl = self.client.get(
            reverse('performance_viewer') + '?year=2026&quarter=all&area=internationalization',
        )
        self.assertEqual(r_sust.status_code, 200)
        self.assertEqual(r_intl.status_code, 200)

        def kpi_total(content):
            m = re.search(rb'viewer-target-ring-inner">(\d+)</div>', content)
            return int(m.group(1)) if m else None

        total_sust = kpi_total(r_sust.content)
        total_intl = kpi_total(r_intl.content)
        self.assertIsNotNone(total_sust)
        self.assertIsNotNone(total_intl)
        self.assertGreater(total_sust, 0)
        self.assertGreater(total_intl, 0)
        self.assertNotEqual(total_sust, total_intl)
        self.assertContains(r_intl, 'Internationalization')
        self.assertContains(r_sust, 'Sustainability')

    def test_performance_viewer_all_development_areas(self):
        call_command('seed_sample_data', '--reset', verbosity=0)
        self.client.login(username='root', password='pass12345')
        r_all = self.client.get(
            reverse('performance_viewer') + '?year=2026&quarter=all&area=all',
        )
        r_one = self.client.get(
            reverse('performance_viewer') + '?year=2026&quarter=all&area=sustainability',
        )
        self.assertEqual(r_all.status_code, 200)
        self.assertEqual(r_one.status_code, 200)
        self.assertContains(r_all, 'All development areas')

        def kpi_total(content):
            m = re.search(rb'viewer-target-ring-inner">(\d+)</div>', content)
            return int(m.group(1)) if m else None

        total_all = kpi_total(r_all.content)
        total_one = kpi_total(r_one.content)
        self.assertIsNotNone(total_all)
        self.assertIsNotNone(total_one)
        self.assertGreaterEqual(total_all, total_one)

    def test_performance_viewer_superuser_bare_url_defaults_to_full_overview(self):
        call_command('seed_sample_data', '--reset', verbosity=0)
        self.client.login(username='root', password='pass12345')
        resp = self.client.get(reverse('performance_viewer'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'All quarters')
        self.assertContains(resp, 'All development areas')
        self.assertContains(resp, 'All offices (combined)')


class SearchSuggestionsHierarchyTests(TestCase):
    """Outcome / Strategy / PAP labels must be searchable (ingest stores office only on PAP rows)."""

    def setUp(self):
        self.client = Client()
        self.office = Office.objects.create(name='SearchOfficeUnique')
        self.user = User.objects.create_user(
            username='search_kpi_user',
            password='pwtest12',
            first_name='SearchOfficeUnique',
        )
        self.outcome = StrategicLevel.objects.create(
            name='SearchOutcome Alpha academic pillar',
            level_type='OUTCOME',
        )
        self.strategy = StrategicLevel.objects.create(
            name='SearchStrategy Beta research line',
            level_type='STRATEGY',
            parent=self.outcome,
        )
        self.pap = StrategicLevel.objects.create(
            name='SearchPAP Gamma program',
            level_type='PAP',
            parent=self.strategy,
            office=self.office,
        )
        self.indicator = Indicator.objects.create(
            pap=self.pap,
            description='Unrelated KPI title for search smoke test',
        )
        self.client.login(username='search_kpi_user', password='pwtest12')

    def test_search_by_outcome_column_text(self):
        resp = self.client.get(reverse('search_suggestions'), {'q': 'academic'})
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, self.indicator.description)
        self.assertNotContains(resp, 'No matching targets')
        self.assertContains(resp, 'target-primary-kicker')
        self.assertContains(resp, self.outcome.name[:100])

    def test_search_by_strategy_column_text(self):
        resp = self.client.get(reverse('search_suggestions'), {'q': 'research'})
        self.assertContains(resp, self.indicator.description)
        self.assertContains(resp, 'target-primary-kicker')
        self.assertContains(resp, self.strategy.name[:100])

    def test_search_by_pap_column_text(self):
        resp = self.client.get(reverse('search_suggestions'), {'q': 'Gamma'})
        self.assertContains(resp, self.indicator.description)
        self.assertContains(resp, self.pap.name[:80])

    def test_search_indicator_text_is_primary_when_query_matches_kpi_only(self):
        resp = self.client.get(reverse('search_suggestions'), {'q': 'smoke test'})
        self.assertContains(resp, 'target-primary-kicker')
        self.assertContains(resp, 'Unrelated KPI title for search smoke test')

    def test_focus_prefers_strategy_when_query_also_in_indicator(self):
        self.assertEqual(
            _search_suggestion_focus(
                'sharedtoken',
                'Indicator sharedtoken text',
                '',
                'Outcome O',
                'Strategy sharedtoken branch',
                'PAP P',
            ),
            'strategy',
        )


class AnnouncementVisibilityTests(TestCase):
    """Announcements: global vs office-targeted visibility."""

    def test_global_active_visible_to_office_linked_user(self):
        Office.objects.create(name='PlanOffice')
        user = User.objects.create_user(username='pdo', password='pwtest12', first_name='PlanOffice')
        Announcement.objects.create(
            title='FY reminder',
            body='Please upload Q4 monitors by Friday.',
            scope=Announcement.SCOPE_GLOBAL,
            is_active=True,
        )
        self.assertEqual(Announcement.objects.visible_for(user).count(), 1)

    def test_office_scoped_visible_only_to_listed_office(self):
        o1 = Office.objects.create(name='Alpha Office')
        o2 = Office.objects.create(name='Beta Office')
        u_alpha = User.objects.create_user(username='a', password='pw', first_name='Alpha Office')
        u_beta = User.objects.create_user(username='b', password='pw', first_name='Beta Office')
        ann = Announcement.objects.create(
            title='Alpha only',
            body='Workshop at Alpha campus.',
            scope=Announcement.SCOPE_OFFICES,
            is_active=True,
        )
        ann.offices.add(o1)
        self.assertTrue(Announcement.objects.visible_for(u_alpha).filter(pk=ann.pk).exists())
        self.assertFalse(Announcement.objects.visible_for(u_beta).filter(pk=ann.pk).exists())

    def test_superuser_sees_active_office_scoped_announcements(self):
        Office.objects.create(name='Gamma Office')
        superuser = User.objects.create_user(username='root', password='pw', is_superuser=True)
        ann = Announcement.objects.create(
            title='Scoped',
            body='X',
            scope=Announcement.SCOPE_OFFICES,
            is_active=True,
        )
        ann.offices.add(Office.objects.get(name='Gamma Office'))
        self.assertTrue(Announcement.objects.visible_for(superuser).filter(pk=ann.pk).exists())


class AnnouncementReadReceiptTests(TestCase):
    """Bell badge reflects unread announcements; mark-read persists."""

    def setUp(self):
        self.client = Client()
        Office.objects.create(name='ReadOffice')
        self.user = User.objects.create_user(username='reader', password='pw', first_name='ReadOffice')
        self.ann = Announcement.objects.create(
            title='Unread post',
            body='Hello',
            scope=Announcement.SCOPE_GLOBAL,
            is_active=True,
        )

    def test_dashboard_shows_badge_then_hides_after_mark_read(self):
        self.client.login(username='reader', password='pw')
        r1 = self.client.get(reverse('performance_viewer'))
        self.assertEqual(r1.status_code, 200)
        self.assertContains(r1, 'class="sopm-announce-nav-badge"')
        r2 = self.client.post(
            reverse('announcement_mark_read'),
            {'ids': str(self.ann.pk)},
        )
        self.assertEqual(r2.status_code, 200)
        self.assertTrue(r2.json().get('ok'))
        r3 = self.client.get(reverse('performance_viewer'))
        self.assertNotContains(r3, 'class="sopm-announce-nav-badge"')

    def test_mark_read_ignores_ids_user_cannot_see(self):
        other_office = Office.objects.create(name='SecretOffice')
        hidden = Announcement.objects.create(
            title='Secret',
            body='X',
            scope=Announcement.SCOPE_OFFICES,
            is_active=True,
        )
        hidden.offices.add(other_office)
        self.client.login(username='reader', password='pw')
        self.client.post(reverse('announcement_mark_read'), {'ids': f'{self.ann.pk},{hidden.pk}'})
        self.assertTrue(AnnouncementRead.objects.filter(user=self.user, announcement=self.ann).exists())
        self.assertFalse(AnnouncementRead.objects.filter(user=self.user, announcement=hidden).exists())


class AnalyticsPillarTooltipLabelTests(TestCase):
    """Tooltip lines must match the hierarchy field that buckets the row, not outcome alone."""

    def test_social_responsibility_prefers_strategy_when_outcome_is_other_pillar(self):
        out = MagicMock()
        out.name = 'Academic Leadership — generic matrix outcome'
        strat = MagicMock()
        strat.name = 'Community extension and outreach services'
        strat.parent = out
        pap = MagicMock()
        pap.name = 'PAP activities'
        pap.parent = strat
        ind = MagicMock()
        ind.description = 'Indicator KPI one'
        ind.pap = pap
        record = MagicMock()
        record.indicator = ind
        label = _analytics_row_tooltip_label_for_pillar(record, 'social_responsibility')
        self.assertIn('extension', label.lower())
        self.assertNotEqual(label, out.name)


class DatabaseBackupDownloadTests(TransactionTestCase):
    """Superusers can download a SQLite snapshot; others are sent to the login page."""

    reset_sequences = True

    def setUp(self):
        self.client = Client()
        self.admin = User.objects.create_user(
            username='backupadmin',
            password='pass12345',
            is_superuser=True,
            is_staff=True,
        )
        User.objects.create_user(username='officestaff', password='pwtest12', first_name='Some Office')

    def test_superuser_gets_sqlite_backup_file(self):
        self.client.login(username='backupadmin', password='pass12345')
        resp = self.client.get(reverse('download_database_backup'))
        self.assertEqual(resp.status_code, 200)
        cd = resp.get('Content-Disposition', '')
        self.assertIn('attachment', cd)
        self.assertIn('.sqlite3', cd)
        self.assertTrue(resp.content.startswith(b'SQLite format 3\x00'))
        self.assertTrue(
            ActivityLog.objects.filter(action='Database backup downloaded').exists()
        )

    def test_office_user_redirected_to_login(self):
        self.client.login(username='officestaff', password='pwtest12')
        resp = self.client.get(reverse('download_database_backup'))
        self.assertEqual(resp.status_code, 302)

    def test_anonymous_redirected_to_login(self):
        resp = self.client.get(reverse('download_database_backup'))
        self.assertEqual(resp.status_code, 302)


class ValidateSqliteBackupPathTests(TestCase):
    def test_rejects_too_small(self):
        fd, p = tempfile.mkstemp(suffix='.sqlite3')
        os.close(fd)
        try:
            Path(p).write_bytes(b'short')
            ok, msg = _validate_sqlite_backup_path(Path(p))
            self.assertFalse(ok)
            self.assertIn('small', msg.lower())
        finally:
            os.unlink(p)

    def test_rejects_wrong_magic(self):
        fd, p = tempfile.mkstemp(suffix='.sqlite3')
        os.close(fd)
        try:
            Path(p).write_bytes(b'x' * 200)
            ok, msg = _validate_sqlite_backup_path(Path(p))
            self.assertFalse(ok)
        finally:
            os.unlink(p)


class RestoreDatabaseBackupViewTests(TransactionTestCase):
    reset_sequences = True

    def setUp(self):
        self.client = Client()
        User.objects.create_user(
            username='restoreadmin',
            password='pass12345',
            is_superuser=True,
            is_staff=True,
        )

    def test_get_restore_page_superuser(self):
        self.client.login(username='restoreadmin', password='pass12345')
        resp = self.client.get(reverse('restore_database_backup'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Restore database from backup')

    def test_get_restore_disabled_when_debug_off(self):
        self.client.login(username='restoreadmin', password='pass12345')
        with override_settings(DEBUG=False, SOPM_ENABLE_DATABASE_RESTORE=False):
            resp = self.client.get(reverse('restore_database_backup'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Restore is turned off')

    def test_post_without_confirm(self):
        self.client.login(username='restoreadmin', password='pass12345')
        upload = SimpleUploadedFile('x.sqlite3', b'x' * 200, content_type='application/octet-stream')
        resp = self.client.post(
            reverse('restore_database_backup'),
            {'confirm': '', 'backup_file': upload},
        )
        self.assertEqual(resp.status_code, 302)

    def test_post_invalid_file_rejected(self):
        self.client.login(username='restoreadmin', password='pass12345')
        upload = SimpleUploadedFile('bad.sqlite3', b'x' * 200, content_type='application/octet-stream')
        resp = self.client.post(
            reverse('restore_database_backup'),
            {'confirm': '1', 'backup_file': upload},
        )
        self.assertEqual(resp.status_code, 302)

    def test_validate_accepts_sqlite_with_django_marker_tables(self):
        """Minimal SQLite with django_migrations + auth_user passes validation."""
        fd, p = tempfile.mkstemp(suffix='.sqlite3')
        os.close(fd)
        try:
            cx = sqlite3.connect(p)
            try:
                cx.execute('CREATE TABLE django_migrations (id integer);')
                cx.execute('CREATE TABLE auth_user (id integer);')
                cx.commit()
            finally:
                cx.close()
            ok, msg = _validate_sqlite_backup_path(Path(p))
            self.assertTrue(ok, msg)
        finally:
            os.unlink(p)
